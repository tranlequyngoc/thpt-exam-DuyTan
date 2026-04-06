# -*- coding: utf-8 -*-
# X-EXAM v5.0 — He thong luyen thi THPT
# Features: Groups, Announcements, Exam Schedule, Source Credits
import os, json, hashlib, secrets, re, sqlite3, random
from datetime import datetime, timedelta, date
from functools import wraps
from flask import Flask, render_template, request, jsonify, redirect, session
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'xexam-duytan-secret-2026-stable')
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
DB = os.path.join(os.path.dirname(__file__), 'data', 'xexam.db')

# Auto-create directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.dirname(DB), exist_ok=True)

# ─── DB ───
DATABASE_URL = os.environ.get('DATABASE_URL', '')
CLOUDINARY_URL = os.environ.get('CLOUDINARY_URL', '')

def upload_to_cloud(file_obj, filename):
    if CLOUDINARY_URL:
        try:
            import cloudinary, cloudinary.uploader
            cloudinary.config(cloudinary_url=CLOUDINARY_URL)
            result = cloudinary.uploader.upload(file_obj, public_id=f'xexam/{filename}', resource_type='auto')
            return result['secure_url']
        except Exception as e:
            print(f"Cloudinary error: {e}")
    fp = os.path.join(app.config['UPLOAD_FOLDER'] if 'app' in dir() else 'uploads', filename)
    if hasattr(file_obj, 'save'): file_obj.save(fp)
    return f'/uploads/{filename}'

class PGCursor:
    def __init__(self, cur):
        self._c = cur; self.lastrowid = None
    def fetchone(self):
        try:
            r = self._c.fetchone()
            return dict(r) if r else None
        except: return None
    def fetchall(self):
        try: return [dict(r) for r in self._c.fetchall()]
        except: return []

class PGConn:
    def __init__(self):
        import psycopg2, psycopg2.extras
        self._pg = psycopg2; self._extras = psycopg2.extras
        self._conn = psycopg2.connect(DATABASE_URL, sslmode='require')
    def execute(self, sql, params=None):
        cur = self._conn.cursor(cursor_factory=self._extras.RealDictCursor)
        sql = sql.replace('?', '%s')
        is_ignore = bool(re.search(r'INSERT\s+OR\s+IGNORE', sql, re.IGNORECASE))
        if is_ignore:
            sql = re.sub(r'INSERT\s+OR\s+IGNORE', 'INSERT', sql, flags=re.IGNORECASE)
            sql = sql.rstrip().rstrip(';') + ' ON CONFLICT DO NOTHING'
        is_insert = sql.strip().upper().startswith('INSERT') and not is_ignore
        if is_insert and 'RETURNING' not in sql.upper():
            sql = sql.rstrip().rstrip(';') + ' RETURNING id'
        try:
            cur.execute(sql, tuple(params) if params else None)
        except self._pg.Error as e:
            self._conn.rollback()
            print(f"PG Error: {e}")
            cur = self._conn.cursor(cursor_factory=self._extras.RealDictCursor)
            return PGCursor(cur)
        wrapper = PGCursor(cur)
        if is_insert:
            try:
                row = cur.fetchone()
                wrapper.lastrowid = row['id'] if row else None
            except: pass
        return wrapper
    def executescript(self, sql):
        sql = re.sub(r'INTEGER\s+PRIMARY\s+KEY\s+AUTOINCREMENT', 'SERIAL PRIMARY KEY', sql, flags=re.IGNORECASE)
        self._conn.autocommit = True
        cur = self._conn.cursor()
        for stmt in sql.split(';'):
            stmt = stmt.strip()
            if not stmt: continue
            try: cur.execute(stmt)
            except Exception as e: print(f"PG skip: {str(e)[:80]}")
        self._conn.autocommit = False
    def commit(self): self._conn.commit()
    def close(self): self._conn.close()

def get_db():
    if DATABASE_URL:
        return PGConn()
    c = sqlite3.connect(DB)
    c.row_factory = sqlite3.Row
    c.execute("PRAGMA foreign_keys=ON")
    return c

def init_db():
    c = get_db()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            fullname TEXT NOT NULL,
            role TEXT DEFAULT 'student',
            avatar TEXT DEFAULT '🎓',
            is_approved INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    CREATE TABLE IF NOT EXISTS exams (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            subject TEXT NOT NULL,
            description TEXT DEFAULT '',
            source TEXT DEFAULT '',
            time_limit INTEGER DEFAULT 90,
            total_score REAL DEFAULT 10.0,
            is_active INTEGER DEFAULT 1,
            is_open INTEGER DEFAULT 1,
            open_at TEXT DEFAULT NULL,
            close_at TEXT DEFAULT NULL,
            max_attempts INTEGER DEFAULT 999,
            shuffle_questions INTEGER DEFAULT 0,
            teacher_approved INTEGER DEFAULT 1,
            created_by INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            exam_id INTEGER,
            question_number INTEGER,
            type TEXT DEFAULT 'multiple_choice',
            content TEXT NOT NULL,
            option_a TEXT DEFAULT '',
            option_b TEXT DEFAULT '',
            option_c TEXT DEFAULT '',
            option_d TEXT DEFAULT '',
            correct_answer TEXT DEFAULT '',
            score REAL DEFAULT 0.25,
            explanation TEXT DEFAULT ''
    );
    CREATE TABLE IF NOT EXISTS submissions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id INTEGER,
            exam_id INTEGER,
            answers TEXT DEFAULT '{}',
            score REAL DEFAULT 0,
            total_correct INTEGER DEFAULT 0,
            total_wrong INTEGER DEFAULT 0,
            total_blank INTEGER DEFAULT 0,
            time_spent INTEGER DEFAULT 0,
            teacher_comment TEXT DEFAULT '',
            violations INTEGER DEFAULT 0,
            submitted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    CREATE TABLE IF NOT EXISTS groups (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            subject TEXT DEFAULT '',
            description TEXT DEFAULT '',
            invite_code TEXT UNIQUE NOT NULL,
            created_by INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    CREATE TABLE IF NOT EXISTS group_members (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            group_id INTEGER,
            user_id INTEGER,
            joined_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(group_id, user_id)
    );
    CREATE TABLE IF NOT EXISTS announcements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            group_id INTEGER,
            author_id INTEGER,
            title TEXT NOT NULL,
            content TEXT NOT NULL,
            exam_id INTEGER DEFAULT NULL,
            pin_exam_at TEXT DEFAULT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            title TEXT NOT NULL,
            message TEXT DEFAULT '',
            link TEXT DEFAULT '',
            is_read INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    """)
    # Migrate exams table if missing columns
    if DATABASE_URL:
        existing = [r['column_name'] for r in c.execute("SELECT column_name FROM information_schema.columns WHERE table_name='exams'").fetchall()]
    else:
        existing = [r[1] for r in c.execute("PRAGMA table_info(exams)").fetchall()]
    for col, defval in [('source',"TEXT DEFAULT ''"), ('is_open',"INTEGER DEFAULT 1"),
                             ('open_at',"TEXT DEFAULT NULL"), ('close_at',"TEXT DEFAULT NULL"),
                             ('max_attempts',"INTEGER DEFAULT 999"), ('shuffle_questions',"INTEGER DEFAULT 0"),
                             ('teacher_approved',"INTEGER DEFAULT 1")]:
            if col not in existing:
                try: c.execute(f"ALTER TABLE exams ADD COLUMN {col} {defval}")
                except: pass
    # Migrate submissions
    if DATABASE_URL:
        existing2 = [r['column_name'] for r in c.execute("SELECT column_name FROM information_schema.columns WHERE table_name='submissions'").fetchall()]
    else:
        existing2 = [r[1] for r in c.execute("PRAGMA table_info(submissions)").fetchall()]
    if 'teacher_comment' not in existing2:
            try: c.execute("ALTER TABLE submissions ADD COLUMN teacher_comment TEXT DEFAULT ''")
            except: pass
    # Migrate users
    if DATABASE_URL:
        existing3 = [r['column_name'] for r in c.execute("SELECT column_name FROM information_schema.columns WHERE table_name='users'").fetchall()]
    else:
        existing3 = [r[1] for r in c.execute("PRAGMA table_info(users)").fetchall()]
    for col, defval in [('avatar',"TEXT DEFAULT '🎓'"), ('is_approved',"INTEGER DEFAULT 1")]:
            if col not in existing3:
                try: c.execute(f"ALTER TABLE users ADD COLUMN {col} {defval}")
                except: pass
    # Migrate submissions - add question_scores
    if DATABASE_URL:
        existing4 = [r['column_name'] for r in c.execute("SELECT column_name FROM information_schema.columns WHERE table_name='submissions'").fetchall()]
    else:
        existing4 = [r[1] for r in c.execute("PRAGMA table_info(submissions)").fetchall()]
    if 'question_scores' not in existing4:
        try: c.execute("ALTER TABLE submissions ADD COLUMN question_scores TEXT DEFAULT '{}'")
        except: pass
    if 'violations' not in existing4:
        try: c.execute("ALTER TABLE submissions ADD COLUMN violations INTEGER DEFAULT 0")
        except: pass

    pw = hashlib.sha256('admin123'.encode()).hexdigest()
    pw2 = hashlib.sha256('123456'.encode()).hexdigest()
    pw_super = hashlib.sha256('DuyTan@2026'.encode()).hexdigest()
    try:
            c.execute("INSERT OR IGNORE INTO users(username,password_hash,fullname,role,avatar) VALUES(?,?,?,?,?)",
                      ('admin', pw, 'Quản trị viên', 'admin', '👑'))
            c.execute("INSERT OR IGNORE INTO users(username,password_hash,fullname,role,avatar) VALUES(?,?,?,?,?)",
                      ('teacher', pw, 'Giáo viên Demo', 'teacher', '👨‍🏫'))
            c.execute("INSERT OR IGNORE INTO users(username,password_hash,fullname,role,avatar) VALUES(?,?,?,?,?)",
                      ('student', pw2, 'Học sinh Demo', 'student', '🎓'))
            c.execute("INSERT OR IGNORE INTO users(username,password_hash,fullname,role,avatar) VALUES(?,?,?,?,?)",
                      ('superadmin', pw_super, 'Super Admin DuyTan', 'admin', '🛡️'))
            c.commit()
    except: pass
    c.close()

# ─── AUTH ───
def login_req(f):
    @wraps(f)
    def w(*a, **k):
            if 'user_id' not in session: return redirect('/login')
            return f(*a, **k)
    return w

def teacher_req(f):
    @wraps(f)
    def w(*a, **k):
            if 'user_id' not in session: return redirect('/login')
            if session.get('role') not in ('teacher', 'admin'):
                return jsonify({'error': 'Không có quyền'}), 403
            return f(*a, **k)
    return w

def admin_req(f):
    @wraps(f)
    def w(*a, **k):
            if 'user_id' not in session: return redirect('/login')
            if session.get('role') != 'admin':
                return jsonify({'error': 'Admin only'}), 403
            return f(*a, **k)
    return w

def get_user(): return {'id': session.get('user_id'), 'role': session.get('role'), 'fullname': session.get('fullname'), 'username': session.get('username')}

# ─── PAGES ───
@app.route('/')
def index():
    return redirect('/dashboard' if 'user_id' in session else '/login')

@app.route('/login')
def login_page(): return render_template('login.html')

@app.route('/dashboard')
@login_req
def dashboard(): return render_template('dashboard.html')

@app.route('/exam/<int:eid>')
@login_req
def exam_page(eid): return render_template('exam.html', exam_id=eid)

@app.route('/result/<int:sid>')
@login_req
def result_page(sid): return render_template('result.html', submission_id=sid)

@app.route('/create-exam')
@teacher_req
def create_exam_page(): return render_template('create_exam.html')

@app.route('/stats')
@login_req
def stats_page(): return render_template('stats.html')

@app.route('/leaderboard')
@login_req
def leaderboard_page(): return render_template('leaderboard.html')

@app.route('/teacher-review')
@teacher_req
def teacher_review_page(): return render_template('teacher_review.html')

@app.route('/chat')
@login_req
def chat_page(): return render_template('chat.html')

@app.route('/groups')
@login_req
def groups_page(): return render_template('groups.html')

@app.route('/group/<int:gid>')
@login_req
def group_detail_page(gid): return render_template('group_detail.html', group_id=gid)

@app.route('/admin')
@admin_req
def admin_page(): return render_template('admin.html')

@app.route('/teacher-stats')
@teacher_req
def teacher_stats_page(): return render_template('teacher_stats.html')

@app.route('/import-students')
@teacher_req
def import_students_page(): return render_template('import_students.html')

@app.route('/static/sw.js')
def sw_js():
    from flask import send_from_directory
    return send_from_directory('static', 'sw.js', mimetype='application/javascript')

# ─── AUTH API ───
@app.route('/api/login', methods=['POST'])
def api_login():
    d = request.json or {}
    pw = hashlib.sha256(d.get('password','').encode()).hexdigest()
    c = get_db()
    u = c.execute("SELECT * FROM users WHERE username=? AND password_hash=?",
                      (d.get('username','').strip(), pw)).fetchone()
    c.close()
    if u:
            session['user_id'] = u['id']
            session['username'] = u['username']
            session['fullname'] = u['fullname']
            session['role'] = u['role']
            return jsonify({'success': True, 'user': dict(u)})
    return jsonify({'success': False, 'error': 'Sai tài khoản hoặc mật khẩu'}), 401

@app.route('/api/register', methods=['POST'])
def api_register():
    d = request.json or {}
    pw = hashlib.sha256(d.get('password','').encode()).hexdigest()
    avatars = {'student':'🎓','teacher':'👨‍🏫','admin':'👑'}
    av = avatars.get(d.get('role','student'),'🎓')
    c = get_db()
    try:
            c.execute("INSERT INTO users(username,password_hash,fullname,role,avatar) VALUES(?,?,?,?,?)",
                      (d['username'].strip(), pw, d['fullname'].strip(), d.get('role','student'), av))
            c.commit(); c.close()
            return jsonify({'success': True})
    except:
            c.close()
            return jsonify({'error': 'Tên đăng nhập đã tồn tại'}), 400

@app.route('/api/logout')
def api_logout():
    session.clear()
    return redirect('/login')

# ─── EXAM APIs ───
@app.route('/api/exams')
@login_req
def api_exams():
    c = get_db()
    uid = session['user_id']
    now = datetime.now().isoformat()
    exams = c.execute("""
            SELECT e.*, u.fullname as creator_name, u.avatar as creator_avatar,
            (SELECT COUNT(*) FROM questions WHERE exam_id=e.id) as question_count,
            (SELECT COUNT(*) FROM submissions WHERE exam_id=e.id AND student_id=?) as my_attempts,
            (SELECT MAX(score) FROM submissions WHERE exam_id=e.id AND student_id=?) as my_best
            FROM exams e LEFT JOIN users u ON e.created_by=u.id
            WHERE e.is_active=1 AND (e.teacher_approved=1 OR e.created_by=?) ORDER BY e.created_at DESC
    """, (uid, uid, uid)).fetchall()
    result = []
    for e in exams:
            ed = dict(e)
            # Check schedule
            now_str = datetime.now().isoformat()
            open_ok = True
            if ed.get('open_at') and now_str < ed['open_at']:
                open_ok = False
            if ed.get('close_at') and now_str > ed['close_at']:
                open_ok = False
            ed['schedule_open'] = open_ok and bool(ed.get('is_open', 1))
            # Check if can attempt more
            attempts = ed.get('my_attempts', 0)
            max_attempts = ed.get('max_attempts', 999)
            ed['can_attempt'] = attempts < max_attempts
            ed['attempts_left'] = max(0, max_attempts - attempts)
            result.append(ed)
    c.close()
    return jsonify(result)

@app.route('/api/exam/<int:eid>')
@login_req
def api_exam(eid):
    c = get_db()
    exam = c.execute("SELECT * FROM exams WHERE id=?", (eid,)).fetchone()
    if not exam: c.close(); return jsonify({'error': 'Không tìm thấy'}), 404
    # Check max attempts
    attempts = c.execute("SELECT COUNT(*) as c FROM submissions WHERE exam_id=? AND student_id=?", (eid, session['user_id'])).fetchone()['c']
    max_attempts = exam.get('max_attempts', 999)
    if attempts >= max_attempts:
        c.close()
        return jsonify({'error': f'Bạn đã thi {attempts} lần. Tối đa {max_attempts} lần.'}), 403
    qs = c.execute("SELECT id,question_number,type,content,option_a,option_b,option_c,option_d,score FROM questions WHERE exam_id=? ORDER BY question_number", (eid,)).fetchall()
    # Shuffle if enabled
    qs_list = [dict(q) for q in qs]
    if exam.get('shuffle_questions', 0):
        random.shuffle(qs_list)
    c.close()
    return jsonify({'exam': dict(exam), 'questions': qs_list})

@app.route('/api/exam/create', methods=['POST'])
@teacher_req
def api_create_exam():
    d = request.json or {}
    c = get_db()
    user = c.execute("SELECT is_approved FROM users WHERE id=?", (session['user_id'],)).fetchone()
    teacher_approved = user['is_approved'] if user else 0
    cur = c.execute("""INSERT INTO exams(title,subject,description,source,time_limit,total_score,created_by,is_open,open_at,close_at,max_attempts,shuffle_questions,teacher_approved)
                           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (d['title'], d['subject'], d.get('description',''), d.get('source',''),
             d.get('time_limit',90), d.get('total_score',10), session['user_id'],
             d.get('is_open',1), d.get('open_at') or None, d.get('close_at') or None,
             d.get('max_attempts',999), d.get('shuffle_questions',0), teacher_approved))
    eid = cur.lastrowid
    for i, q in enumerate(d.get('questions',[]), 1):
            c.execute("""INSERT INTO questions(exam_id,question_number,type,content,option_a,option_b,option_c,option_d,correct_answer,score,explanation)
                         VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                (eid, i, q.get('type','multiple_choice'), q.get('content',''),
                 q.get('option_a',''), q.get('option_b',''), q.get('option_c',''), q.get('option_d',''),
                 q.get('correct_answer',''), q.get('score',0.25), q.get('explanation','')))
    # Notify students in teacher's groups
    try:
        groups = c.execute("SELECT group_id FROM group_members WHERE user_id=?", (session['user_id'],)).fetchall()
        for g in groups:
            members = c.execute("SELECT user_id FROM group_members WHERE group_id=? AND user_id!=?", (g['group_id'], session['user_id'])).fetchall()
            for m in members:
                c.execute("INSERT INTO notifications(user_id,title,message,link) VALUES(?,?,?,?)",
                    (m['user_id'], 'De thi moi!', f'Giao vien vua dang de "{d["title"]}" - Mon {d["subject"]}', f'/exam/{eid}'))
    except: pass
    c.commit(); c.close()
    return jsonify({'success': True, 'exam_id': eid})

@app.route('/api/exam/toggle/<int:eid>', methods=['POST'])
@teacher_req
def api_toggle_exam(eid):
    c = get_db()
    exam = c.execute("SELECT is_open FROM exams WHERE id=?", (eid,)).fetchone()
    if not exam: c.close(); return jsonify({'error': 'Không tìm thấy'}), 404
    new_val = 0 if exam['is_open'] else 1
    c.execute("UPDATE exams SET is_open=? WHERE id=?", (new_val, eid))
    c.commit(); c.close()
    return jsonify({'success': True, 'is_open': new_val})

@app.route('/api/exam/delete/<int:eid>', methods=['DELETE'])
@teacher_req
def api_delete_exam(eid):
    c = get_db()
    c.execute("DELETE FROM questions WHERE exam_id=?", (eid,))
    c.execute("DELETE FROM submissions WHERE exam_id=?", (eid,))
    c.execute("DELETE FROM exams WHERE id=?", (eid,))
    c.commit(); c.close()
    return jsonify({'success': True})

# ─── SUBMIT ───
@app.route('/api/submit', methods=['POST'])
@login_req
def api_submit():
    d = request.json or {}
    eid = d['exam_id']
    ans = d.get('answers', {})
    c = get_db()
    
    # Check exam is open
    exam = c.execute("SELECT is_open,open_at,close_at FROM exams WHERE id=?", (eid,)).fetchone()
    if exam:
            now_str = datetime.now().isoformat()
            if not exam['is_open']:
                c.close(); return jsonify({'error': 'Bài kiểm tra đã đóng!'}), 403
            if exam['open_at'] and now_str < exam['open_at']:
                c.close(); return jsonify({'error': 'Bài kiểm tra chưa mở!'}), 403
            if exam['close_at'] and now_str > exam['close_at']:
                c.close(); return jsonify({'error': 'Bài kiểm tra đã hết hạn!'}), 403
    
    qs = c.execute("SELECT id,correct_answer,score,type FROM questions WHERE exam_id=?", (eid,)).fetchall()
    sc = 0.0; cor = 0; wr = 0; bl = 0
    for q in qs:
            sa = str(ans.get(str(q['id']), '')).strip()
            ca = str(q['correct_answer'] or '').strip()
            qt = q['type'] or 'multiple_choice'
            if qt == 'essay':
                if sa: sc += q['score'] * 0.5; cor += 1
                else: bl += 1
            elif qt == 'true_false':
                if not sa: bl += 1
                else:
                    total_sub = max(len(ca), 1)
                    correct_sub = sum(1 for i in range(min(len(sa), len(ca))) if sa[i].upper() == ca[i].upper())
                    if correct_sub == total_sub: sc += q['score']; cor += 1
                    elif correct_sub > 0: sc += q['score'] * correct_sub / total_sub; cor += 1
                    else: wr += 1
            else:
                if not sa: bl += 1
                elif sa.upper() == ca.upper(): sc += q['score']; cor += 1
                else: wr += 1
    
    cur = c.execute("INSERT INTO submissions(student_id,exam_id,answers,score,total_correct,total_wrong,total_blank,time_spent,violations) VALUES(?,?,?,?,?,?,?,?,?)",
            (session['user_id'], eid, json.dumps(ans), round(sc, 2), cor, wr, bl, d.get('time_spent', 0), d.get('violations', 0)))
    sid = cur.lastrowid
    c.commit(); c.close()
    return jsonify({'success': True, 'submission_id': sid, 'score': round(sc,2), 'correct': cor, 'wrong': wr, 'blank': bl})

@app.route('/api/result/<int:sid>')
@login_req
def api_result(sid):
    c = get_db()
    sub = c.execute("SELECT * FROM submissions WHERE id=?", (sid,)).fetchone()
    if not sub: c.close(); return jsonify({'error': 'Không tìm thấy'}), 404
    exam = c.execute("SELECT * FROM exams WHERE id=?", (sub['exam_id'],)).fetchone()
    qs = c.execute("SELECT * FROM questions WHERE exam_id=? ORDER BY question_number", (sub['exam_id'],)).fetchall()
    c.close()
    return jsonify({'submission': dict(sub), 'exam': dict(exam), 'questions': [dict(q) for q in qs],
                        'student_answers': json.loads(sub['answers'] or '{}')})

# ─── STATS ───
@app.route('/api/stats')
@login_req
def api_stats():
    uid = session['user_id']; c = get_db()
    te = c.execute("SELECT COUNT(DISTINCT exam_id) as c FROM submissions WHERE student_id=?", (uid,)).fetchone()['c']
    ts = c.execute("SELECT COUNT(*) as c FROM submissions WHERE student_id=?", (uid,)).fetchone()['c']
    av = c.execute("SELECT AVG(score) as a FROM submissions WHERE student_id=?", (uid,)).fetchone()['a'] or 0
    bs = c.execute("SELECT e.subject, AVG(s.score) as avg_score, COUNT(*) as attempts, MAX(s.score) as best_score FROM submissions s JOIN exams e ON s.exam_id=e.id WHERE s.student_id=? GROUP BY e.subject", (uid,)).fetchall()
    rc = c.execute("SELECT s.*,e.title,e.subject,e.total_score FROM submissions s JOIN exams e ON s.exam_id=e.id WHERE s.student_id=? ORDER BY s.submitted_at DESC LIMIT 20", (uid,)).fetchall()
    c.close()
    return jsonify({'total_exams': te, 'total_submissions': ts, 'avg_score': round(av,2),
                        'by_subject': [dict(s) for s in bs], 'recent': [dict(r) for r in rc]})

@app.route('/api/my-rank')
@login_req
def api_my_rank():
    c = get_db()
    uid = session['user_id']
    total = c.execute("SELECT COUNT(*) as c FROM submissions WHERE student_id=?", (uid,)).fetchone()['c']
    c.close()
    rank = get_rank(total)
    rank['meme'] = random.choice(MASCOT_IDLE)
    return jsonify(rank)

@app.route('/api/streak')
@login_req
def api_streak():
    uid = session['user_id']; c = get_db()
    dates = c.execute("SELECT DISTINCT DATE(submitted_at) as d FROM submissions WHERE student_id=? ORDER BY d DESC", (uid,)).fetchall()
    c.close()
    date_list = [d['d'] for d in dates if d['d']]
    today_str = date.today().isoformat()
    today_done = today_str in date_list
    streak = 0
    check = date.today() if today_done else date.today() - timedelta(days=1)
    for i in range(400):
            if (check - timedelta(days=i)).isoformat() in date_list: streak += 1
            else: break
    longest = 0; cur_s = 0
    for i, ds in enumerate(sorted(date_list)):
            if i == 0: cur_s = 1
            else:
                prev = datetime.strptime(sorted(date_list)[i-1],'%Y-%m-%d').date()
                curr = datetime.strptime(ds,'%Y-%m-%d').date()
                cur_s = cur_s+1 if (curr-prev).days==1 else 1
            longest = max(longest, cur_s)
    STREAK_BADGES = [(3,'Chăm học cấp 1','📘'),(7,'Máy giải đề','🤖'),(14,'Quái vật ôn thi','👾'),(30,'Siêu nhân học tập','🦸'),(60,'Huyền thoại','🐉')]
    badges = [{'name':n,'icon':ic,'days':d} for d,n,ic in STREAK_BADGES if longest>=d]
    return jsonify({'streak': streak, 'longest': longest, 'badges': badges, 'today_done': today_done, 'total_days': len(date_list)})

# ─── LEADERBOARD ───
@app.route('/api/leaderboard')
@login_req
def api_leaderboard():
    c = get_db()
    king_count = c.execute("SELECT u.id,u.fullname,u.avatar,COUNT(*) as total_subs,AVG(s.score) as avg_score,MAX(s.score) as best FROM submissions s JOIN users u ON s.student_id=u.id GROUP BY u.id ORDER BY total_subs DESC LIMIT 20").fetchall()
    king_score = c.execute("SELECT u.id,u.fullname,u.avatar,COUNT(*) as total_subs,AVG(s.score) as avg_score,MAX(s.score) as best FROM submissions s JOIN users u ON s.student_id=u.id GROUP BY u.id ORDER BY avg_score DESC LIMIT 20").fetchall()
    king_elite = c.execute("SELECT u.id,u.fullname,u.avatar,COUNT(*) as total_subs,AVG(s.score) as avg_score,MAX(s.score) as best,(COUNT(*)*0.3+AVG(s.score)*0.7) as elite FROM submissions s JOIN users u ON s.student_id=u.id GROUP BY u.id ORDER BY elite DESC LIMIT 20").fetchall()
    # Teacher rank by exams created
    teacher_rank = c.execute("SELECT u.id,u.fullname,u.avatar,COUNT(e.id) as total_exams FROM users u LEFT JOIN exams e ON e.created_by=u.id WHERE u.role IN('teacher','admin') GROUP BY u.id ORDER BY total_exams DESC LIMIT 20").fetchall()
    c.close()
    return jsonify({'king_count':[dict(r) for r in king_count],'king_score':[dict(r) for r in king_score],
                        'king_elite':[dict(r) for r in king_elite],'teacher_rank':[dict(r) for r in teacher_rank]})

# ─── GROUPS ───
@app.route('/api/groups')
@login_req
def api_groups():
    uid = session['user_id']; c = get_db()
    if session['role'] in ('teacher','admin'):
            groups = c.execute("SELECT g.*,(SELECT COUNT(*) FROM group_members WHERE group_id=g.id) as member_count FROM groups g WHERE g.created_by=? ORDER BY g.created_at DESC", (uid,)).fetchall()
    else:
            groups = c.execute("SELECT g.*,(SELECT COUNT(*) FROM group_members WHERE group_id=g.id) as member_count FROM groups g JOIN group_members gm ON g.id=gm.group_id WHERE gm.user_id=? ORDER BY g.created_at DESC", (uid,)).fetchall()
    c.close()
    return jsonify([dict(g) for g in groups])

@app.route('/api/group/create', methods=['POST'])
@teacher_req
def api_create_group():
    d = request.json or {}
    code = secrets.token_urlsafe(6).upper()
    c = get_db()
    cur = c.execute("INSERT INTO groups(name,subject,description,invite_code,created_by) VALUES(?,?,?,?,?)",
            (d['name'], d.get('subject',''), d.get('description',''), code, session['user_id']))
    gid = cur.lastrowid
    c.execute("INSERT OR IGNORE INTO group_members(group_id,user_id) VALUES(?,?)", (gid, session['user_id']))
    c.commit(); c.close()
    return jsonify({'success': True, 'group_id': gid, 'invite_code': code})

@app.route('/api/group/join', methods=['POST'])
@login_req
def api_join_group():
    d = request.json or {}
    code = d.get('code','').strip().upper()
    c = get_db()
    group = c.execute("SELECT * FROM groups WHERE invite_code=?", (code,)).fetchone()
    if not group: c.close(); return jsonify({'error': 'Mã mời không hợp lệ!'}), 404
    try:
            c.execute("INSERT INTO group_members(group_id,user_id) VALUES(?,?)", (group['id'], session['user_id']))
            c.commit()
    except: pass
    c.close()
    return jsonify({'success': True, 'group_name': group['name']})

@app.route('/api/group/<int:gid>')
@login_req
def api_group_detail(gid):
    c = get_db()
    group = c.execute("SELECT * FROM groups WHERE id=?", (gid,)).fetchone()
    if not group: c.close(); return jsonify({'error': 'Không tìm thấy nhóm'}), 404
    members = c.execute("SELECT u.id,u.fullname,u.role,u.avatar FROM group_members gm JOIN users u ON gm.user_id=u.id WHERE gm.group_id=?", (gid,)).fetchall()
    announcements = c.execute("SELECT a.*,u.fullname as author_name,u.avatar as author_avatar FROM announcements a JOIN users u ON a.author_id=u.id WHERE a.group_id=? ORDER BY a.created_at DESC", (gid,)).fetchall()
    c.close()
    return jsonify({'group': dict(group), 'members': [dict(m) for m in members], 'announcements': [dict(a) for a in announcements]})

@app.route('/api/group/<int:gid>/announce', methods=['POST'])
@teacher_req
def api_announce(gid):
    d = request.json or {}
    c = get_db()
    c.execute("INSERT INTO announcements(group_id,author_id,title,content,exam_id,pin_exam_at) VALUES(?,?,?,?,?,?)",
            (gid, session['user_id'], d['title'], d['content'], d.get('exam_id') or None, d.get('pin_exam_at') or None))
    c.commit(); c.close()
    return jsonify({'success': True})

# ─── TEACHER REVIEW ───
@app.route('/api/teacher/submissions')
@teacher_req
def api_teacher_subs():
    c = get_db()
    subs = c.execute("SELECT s.*,u.fullname as student_name,u.username,e.title as exam_title,e.subject,e.total_score FROM submissions s JOIN users u ON s.student_id=u.id JOIN exams e ON s.exam_id=e.id ORDER BY s.submitted_at DESC LIMIT 100").fetchall()
    c.close()
    return jsonify([dict(s) for s in subs])

@app.route('/api/teacher/comment', methods=['POST'])
@teacher_req
def api_comment():
    d = request.json or {}
    c = get_db()
    c.execute("UPDATE submissions SET teacher_comment=? WHERE id=?", (d.get('comment',''), d['submission_id']))
    c.commit(); c.close()
    return jsonify({'success': True})

# ─── QUICK EXAM ───
@app.route('/api/quick-exam', methods=['POST'])
@login_req
def api_quick_exam():
    d = request.json or {}
    subject = d.get('subject',''); num = d.get('num', 10)
    c = get_db()
    if subject:
            qs = c.execute("SELECT q.* FROM questions q JOIN exams e ON q.exam_id=e.id WHERE e.subject=? ORDER BY RANDOM() LIMIT ?", (subject, num)).fetchall()
    else:
            qs = c.execute("SELECT * FROM questions ORDER BY RANDOM() LIMIT ?", (num,)).fetchall()
    if not qs: c.close(); return jsonify({'error': 'Chưa có câu hỏi nào!'}), 400
    title = f"Luyện nhanh {subject or 'Tổng hợp'} — {datetime.now().strftime('%H:%M %d/%m')}"
    cur = c.execute("INSERT INTO exams(title,subject,description,time_limit,total_score,created_by) VALUES(?,?,?,?,?,?)",
            (title, subject or 'Tổng hợp', 'Đề luyện nhanh tự động', 5, 10, session['user_id']))
    eid = cur.lastrowid; spp = round(10.0/len(qs), 2)
    for i, q in enumerate(qs, 1):
            c.execute("INSERT INTO questions(exam_id,question_number,type,content,option_a,option_b,option_c,option_d,correct_answer,score) VALUES(?,?,?,?,?,?,?,?,?,?)",
                (eid, i, q['type'], q['content'], q['option_a'], q['option_b'], q['option_c'], q['option_d'], q['correct_answer'], spp))
    c.commit(); c.close()
    return jsonify({'success': True, 'exam_id': eid, 'num_questions': len(qs)})

# ─── UPLOAD ───
@app.route('/api/upload-exam', methods=['POST'])
@teacher_req
def api_upload():
    if 'file' not in request.files: return jsonify({'error': 'Không có file'}), 400
    f = request.files['file']
    fn = secure_filename(f.filename or 'upload')
    orig_ext = (f.filename or '').rsplit('.',1)[-1].lower() if '.' in (f.filename or '') else 'txt'
    if not fn.endswith('.' + orig_ext):
            fn = fn + '.' + orig_ext
    fp = os.path.join(app.config['UPLOAD_FOLDER'], fn)
    f.save(fp)
    text = read_file_content(fp)
    if not text or len(text) < 5:
            return jsonify({'success': False, 'error': f'Không đọc được file .{orig_ext}', 'questions': []})
    qs = parse_questions_smart(text)
    return jsonify({'success': True, 'filename': fn, 'file_type': orig_ext, 'content': text[:8000],
                        'total_chars': len(text), 'questions': qs,
                        'message': f'Đọc được {len(text)} ký tự, tìm thấy {len(qs)} câu hỏi'})

@app.route('/api/ai-auto-answer', methods=['POST'])
@teacher_req
def api_ai_auto():
    d = request.json or {}
    qs = d.get('questions', [])
    if not isinstance(qs, list):
        return jsonify({'error': 'Dữ liệu không hợp lệ'}), 400
    for q in qs:
        qtype = q.get('type', 'multiple_choice')
        if qtype == 'multiple_choice':
            opts = {'A': q.get('option_a',''), 'B': q.get('option_b',''),
                    'C': q.get('option_c',''), 'D': q.get('option_d','')}
            lens = {k: len(v) for k, v in opts.items() if v and v.strip()}
            if lens:
                q['correct_answer'] = max(lens, key=lens.get)
            else:
                q['correct_answer'] = 'A'
        elif qtype == 'true_false':
            # Analyze content for positive/negative keywords
            content = (q.get('content','') or '').lower()
            tf_ans = []
            sub_lines = [l.strip() for l in content.split('\n') if re.match(r'^[abcd]\)', l.strip(), re.IGNORECASE)]
            if not sub_lines:
                sub_lines = [f'y {x}' for x in ['a','b','c','d']]
            for sub in sub_lines[:4]:
                neg_words = ['khong','sai','false','khong phai','chua','khong the','khong co','0','negative']
                if any(w in sub.lower() for w in neg_words):
                    tf_ans.append('S')
                else:
                    tf_ans.append('D')
            while len(tf_ans) < 4:
                tf_ans.append('D')
            q['correct_answer'] = ''.join(tf_ans[:4])
        elif qtype == 'essay':
            q['correct_answer'] = ''
    return jsonify({'success': True, 'questions': qs,
                    'warning': 'AI chi goi y — Giao vien can kiem tra lai tat ca!'})

# ─── CHAT ───
@app.route('/api/chat', methods=['POST'])
@login_req
def api_chat():
    msg = (request.json or {}).get('message', '').strip()
    if not msg: return jsonify({'reply': 'Em chưa nhập câu hỏi!'})
    return jsonify({'reply': ai_reply(msg)})

# ─── MASCOT ───
@app.route('/api/mascot')
@login_req
def api_mascot():
    action = request.args.get('action', 'idle')
    msgs = {'correct': MASCOT_CORRECT, 'wrong': MASCOT_WRONG, 'idle': MASCOT_IDLE}
    pool = msgs.get(action, MASCOT_IDLE)
    if action == 'submit':
            pct = float(request.args.get('score',0)) / max(float(request.args.get('total',10)),1) * 100
            cat = 'god' if pct>=95 else 'great' if pct>=80 else 'good' if pct>=60 else 'ok' if pct>=40 else 'low'
            pool = MASCOT_SUBMIT.get(cat, MASCOT_IDLE)
    return jsonify({'message': random.choice(pool), 'name': 'Ôn Thi Chan'})

@app.route('/api/result-meme')
@login_req
def api_meme():
    pct = float(request.args.get('score',0)) / max(float(request.args.get('total',10)),1) * 100
    cat = 'god' if pct>=95 else 'great' if pct>=80 else 'good' if pct>=60 else 'ok' if pct>=40 else 'low'
    return jsonify({'meme': random.choice(RESULT_MEMES[cat]), 'category': cat, 'percentage': round(pct,1)})

# ─── ADMIN APIs ───
@app.route('/api/admin/overview')
@admin_req
def api_admin_overview():
    c = get_db()
    tu = c.execute("SELECT COUNT(*) as c FROM users").fetchone()['c']
    ts = c.execute("SELECT COUNT(*) as c FROM users WHERE role='student'").fetchone()['c']
    tt = c.execute("SELECT COUNT(*) as c FROM users WHERE role='teacher'").fetchone()['c']
    ta = c.execute("SELECT COUNT(*) as c FROM users WHERE role='admin'").fetchone()['c']
    te = c.execute("SELECT COUNT(*) as c FROM exams").fetchone()['c']
    tsub = c.execute("SELECT COUNT(*) as c FROM submissions").fetchone()['c']
    tg = c.execute("SELECT COUNT(*) as c FROM groups").fetchone()['c']
    c.close()
    return jsonify({'total_users':tu,'total_students':ts,'total_teachers':tt,'total_admins':ta,'total_exams':te,'total_submissions':tsub,'total_groups':tg})

@app.route('/api/admin/users')
@admin_req
def api_admin_users():
    c = get_db()
    users = c.execute("""SELECT u.id,u.username,u.fullname,u.role,u.avatar,u.is_approved,u.created_at,
        (SELECT COUNT(*) FROM submissions WHERE student_id=u.id) as total_subs,
        (SELECT COUNT(*) FROM exams WHERE created_by=u.id) as total_exams
        FROM users u ORDER BY u.created_at DESC""").fetchall()
    c.close()
    return jsonify([dict(u) for u in users])

@app.route('/api/admin/delete-user/<int:uid>', methods=['DELETE'])
@admin_req
def api_admin_delete_user(uid):
    if uid == session['user_id']: return jsonify({'error':'Khong the xoa chinh minh'}),400
    c = get_db()
    u = c.execute("SELECT username FROM users WHERE id=?", (uid,)).fetchone()
    if not u: c.close(); return jsonify({'error':'User not found'}),404
    c.execute("DELETE FROM submissions WHERE student_id=?", (uid,))
    c.execute("DELETE FROM group_members WHERE user_id=?", (uid,))
    c.execute("DELETE FROM announcements WHERE author_id=?", (uid,))
    c.execute("DELETE FROM exams WHERE created_by=?", (uid,))
    c.execute("DELETE FROM users WHERE id=?", (uid,))
    c.commit(); c.close()
    return jsonify({'success':True})

@app.route('/api/admin/change-role', methods=['POST'])
@admin_req
def api_admin_change_role():
    d = request.json or {}
    uid = d.get('user_id'); new_role = d.get('role')
    if new_role not in ('student','teacher','admin'): return jsonify({'error':'Invalid role'}),400
    if uid == session['user_id']: return jsonify({'error':'Khong the doi role chinh minh'}),400
    avatars = {'student':'🎓','teacher':'👨‍🏫','admin':'👑'}
    c = get_db()
    c.execute("UPDATE users SET role=?,avatar=? WHERE id=?", (new_role, avatars.get(new_role,'🎓'), uid))
    c.commit(); c.close()
    return jsonify({'success':True})

@app.route('/api/admin/approve-teacher/<int:uid>', methods=['POST'])
@admin_req
def api_admin_approve_teacher(uid):
    c = get_db()
    c.execute("UPDATE users SET is_approved=1 WHERE id=? AND role='teacher'", (uid,))
    c.execute("UPDATE exams SET teacher_approved=1 WHERE created_by=?", (uid,))
    c.commit(); c.close()
    return jsonify({'success':True})

@app.route('/api/admin/reject-teacher/<int:uid>', methods=['POST'])
@admin_req
def api_admin_reject_teacher(uid):
    c = get_db()
    c.execute("UPDATE users SET is_approved=0 WHERE id=?", (uid,))
    c.execute("UPDATE exams SET is_active=0 WHERE created_by=?", (uid,))
    c.commit(); c.close()
    return jsonify({'success':True})

@app.route('/api/admin/all-submissions')
@admin_req
def api_admin_all_subs():
    c = get_db()
    subs = c.execute("""SELECT s.*,u.fullname as student_name,u.username,e.title as exam_title,e.subject,e.total_score 
        FROM submissions s JOIN users u ON s.student_id=u.id JOIN exams e ON s.exam_id=e.id 
        ORDER BY s.submitted_at DESC LIMIT 200""").fetchall()
    c.close()
    return jsonify([dict(s) for s in subs])

# ─── IMAGE UPLOAD FOR ESSAY ───
@app.route('/api/upload-answer-image', methods=['POST'])
@login_req
def api_upload_answer_image():
    if 'image' not in request.files: return jsonify({'error':'No image'}),400
    f = request.files['image']
    import time
    ext = (f.filename or 'img.jpg').rsplit('.',1)[-1].lower()
    fn = f"ans_{session['user_id']}_{int(time.time())}.{ext}"
    fp = os.path.join(app.config['UPLOAD_FOLDER'], fn)
    f.save(fp)
    return jsonify({'success':True,'url':f'/uploads/{fn}','filename':fn})

@app.route('/uploads/<filename>')
def serve_upload(filename):
    from flask import send_from_directory
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# ─── NOTIFICATIONS ───
@app.route('/api/notifications')
@login_req
def api_notifications():
    c = get_db()
    uid = session['user_id']
    notifs = c.execute("SELECT * FROM notifications WHERE user_id=? ORDER BY created_at DESC LIMIT 20", (uid,)).fetchall()
    unread = c.execute("SELECT COUNT(*) as c FROM notifications WHERE user_id=? AND is_read=0", (uid,)).fetchone()['c']
    c.close()
    return jsonify({'notifications': [dict(n) for n in notifs], 'unread': unread})

@app.route('/api/notifications/read', methods=['POST'])
@login_req
def api_mark_read():
    c = get_db()
    c.execute("UPDATE notifications SET is_read=1 WHERE user_id=?", (session['user_id'],))
    c.commit(); c.close()
    return jsonify({'success': True})

# ─── IMPORT STUDENTS ───
@app.route('/api/admin/import-students', methods=['POST'])
@teacher_req
def api_import_students():
    if 'file' not in request.files: return jsonify({'error':'No file'}),400
    f = request.files['file']
    results = []; errors = []
    try:
        from openpyxl import load_workbook
        import io, tempfile
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        f.save(tmp.name); tmp.close()
        wb = load_workbook(tmp.name); ws = wb.active
        os.unlink(tmp.name)
        c = get_db()
        for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
            if not row or not row[0]: continue
            fullname = str(row[0] or '').strip()
            username = str(row[1] or '').strip().lower().replace(' ','')
            password = str(row[2] or '123456').strip()
            role = str(row[3] or 'student').strip().lower()
            if role not in ('student','teacher'): role = 'student'
            if not fullname or not username:
                errors.append(f'Dong {i}: Thieu ho ten hoac username'); continue
            pw_hash = hashlib.sha256(password.encode()).hexdigest()
            avatars = {'student':'🎓','teacher':'👨‍🏫'}
            try:
                c.execute("INSERT INTO users(username,password_hash,fullname,role,avatar,is_approved) VALUES(?,?,?,?,?,?)",
                    (username, pw_hash, fullname, role, avatars.get(role,'🎓'), 1))
                results.append({'fullname':fullname,'username':username,'password':password,'role':role})
            except Exception as e:
                errors.append(f'Dong {i}: {username} da ton tai')
        c.commit(); c.close()
    except Exception as e:
        return jsonify({'error':str(e)}), 500
    return jsonify({'success':True,'created':len(results),'errors':errors,'accounts':results})

# ─── RESET PASSWORD ───
@app.route('/api/admin/reset-password', methods=['POST'])
@admin_req
def api_admin_reset_password():
    d = request.json or {}
    uid = d.get('user_id')
    new_pw = d.get('new_password', '')
    if len(new_pw) < 6: return jsonify({'error': 'Mat khau phai co it nhat 6 ky tu'}), 400
    pw_hash = hashlib.sha256(new_pw.encode()).hexdigest()
    c = get_db()
    c.execute("UPDATE users SET password_hash=? WHERE id=?", (pw_hash, uid))
    c.commit(); c.close()
    return jsonify({'success': True})

# ─── EXCEL EXPORT ───
@app.route('/api/teacher/export-excel')
@teacher_req
def api_export_excel():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill, Font
        from flask import send_file
        import io
    except ImportError:
        return jsonify({'error': 'openpyxl chua duoc cai dat'}), 500
    exam_id = request.args.get('exam_id')
    c = get_db()
    if exam_id:
        subs = c.execute("""SELECT s.*,u.fullname as student_name,u.username,e.title as exam_title,e.subject,e.total_score
            FROM submissions s JOIN users u ON s.student_id=u.id JOIN exams e ON s.exam_id=e.id
            WHERE e.id=? ORDER BY s.score DESC""", (exam_id,)).fetchall()
    else:
        subs = c.execute("""SELECT s.*,u.fullname as student_name,u.username,e.title as exam_title,e.subject,e.total_score
            FROM submissions s JOIN users u ON s.student_id=u.id JOIN exams e ON s.exam_id=e.id
            ORDER BY e.title, s.score DESC LIMIT 500""").fetchall()
    c.close()
    wb = Workbook(); ws = wb.active; ws.title = "Bang diem"
    header_fill = PatternFill(start_color='302b63', end_color='302b63', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')
    headers = ['STT','Ho ten','Username','De thi','Mon','Diem','Tong diem','Dung','Sai','Trong','Thoi gian','Ngay nop']
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=i, value=h)
        cell.fill = header_fill; cell.font = header_font
    for i, s in enumerate(subs, 2):
        ws.append([i-1, s['student_name'], s['username'], s['exam_title'], s['subject'],
                   s['score'], s['total_score'], s['total_correct'], s['total_wrong'],
                   s['total_blank'], f"{(s['time_spent']or 0)//60}p{(s['time_spent']or 0)%60}s", s['submitted_at']])
    output = io.BytesIO(); wb.save(output); output.seek(0)
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='bang_diem_xexam.xlsx')

# ─── ESSAY GRADING ───
@app.route('/api/teacher/grade-submission', methods=['POST'])
@teacher_req
def api_grade_submission():
    d = request.json or {}
    sub_id = d.get('submission_id')
    essay_scores = d.get('essay_scores', {})
    comment = d.get('comment', '')
    c = get_db()
    sub = c.execute("SELECT * FROM submissions WHERE id=?", (sub_id,)).fetchone()
    if not sub: c.close(); return jsonify({'error': 'Khong tim thay'}), 404
    qs = c.execute("SELECT * FROM questions WHERE exam_id=?", (sub['exam_id'],)).fetchall()
    answers = json.loads(sub['answers'] or '{}')
    total = 0.0; cor = 0; wr = 0; bl = 0
    for q in qs:
        qid = str(q['id'])
        if q['type'] == 'essay':
            if qid in essay_scores:
                sc = min(float(essay_scores[qid]), float(q['score']))
                total += sc
                if sc > 0: cor += 1
                else: bl += 1
            else:
                sa = str(answers.get(qid, '')).strip()
                if sa: total += float(q['score']) * 0.5; cor += 1
                else: bl += 1
        else:
            sa = str(answers.get(qid, '')).strip()
            ca = str(q['correct_answer'] or '').strip()
            if not sa: bl += 1
            elif sa.upper() == ca.upper(): total += float(q['score']); cor += 1
            else: wr += 1
    c.execute("UPDATE submissions SET score=?,total_correct=?,total_wrong=?,total_blank=?,teacher_comment=?,question_scores=? WHERE id=?",
              (round(total,2), cor, wr, bl, comment, json.dumps(essay_scores), sub_id))
    c.commit(); c.close()
    return jsonify({'success': True, 'new_score': round(total, 2)})

# ─── TEACHER CHART STATS ───
@app.route('/api/teacher/chart-stats')
@teacher_req
def api_teacher_chart_stats():
    c = get_db()
    exam_id = request.args.get('exam_id')
    if exam_id:
        subs = c.execute("SELECT s.score,s.answers,e.total_score FROM submissions s JOIN exams e ON s.exam_id=e.id WHERE s.exam_id=?", (exam_id,)).fetchall()
        qs = c.execute("SELECT id,question_number,content,correct_answer,type FROM questions WHERE exam_id=? ORDER BY question_number", (exam_id,)).fetchall()
        # Question difficulty analysis
        q_stats = []
        for q in qs:
            correct = 0; total_ans = 0
            for sub in subs:
                ans = json.loads(sub['answers'] or '{}')
                sa = str(ans.get(str(q['id']), '')).strip()
                ca = str(q['correct_answer'] or '').strip()
                if sa: total_ans += 1
                if sa.upper() == ca.upper(): correct += 1
            pct = round(correct/max(total_ans,1)*100, 1)
            q_stats.append({'num': q['question_number'], 'correct_pct': pct, 'content': (q['content'] or '')[:50]})
    else:
        subs = c.execute("SELECT s.score,e.total_score FROM submissions s JOIN exams e ON s.exam_id=e.id").fetchall()
        q_stats = []
    # Score distribution
    if subs:
        total_score = subs[0]['total_score'] or 10
        buckets = [0]*5  # 0-2, 2-4, 4-6, 6-8, 8-10
        for s in subs:
            pct = (s['score'] or 0) / total_score * 10
            idx = min(int(pct/2), 4)
            buckets[idx] += 1
    else:
        buckets = [0]*5
    c.close()
    return jsonify({'score_distribution': buckets, 'question_stats': q_stats,
                    'labels': ['0-2','2-4','4-6','6-8','8-10'], 'total': len(subs)})

# ─── CLOUDINARY UPLOAD UPDATE ───
@app.route('/api/upload-answer-image', methods=['POST'])
@login_req
def api_upload_answer_image_v2():
    if 'image' not in request.files: return jsonify({'error':'No image'}),400
    f = request.files['image']
    import time
    ext = (f.filename or 'img.jpg').rsplit('.',1)[-1].lower()
    fn = f"ans_{session['user_id']}_{int(time.time())}.{ext}"
    if CLOUDINARY_URL:
        try:
            import cloudinary, cloudinary.uploader
            cloudinary.config(cloudinary_url=CLOUDINARY_URL)
            result = cloudinary.uploader.upload(f, public_id=f'xexam/answers/{fn}', resource_type='image')
            url = result['secure_url']
            return jsonify({'success':True,'url':url,'filename':fn})
        except Exception as e:
            print(f"Cloudinary error: {e}")
    fp = os.path.join(app.config['UPLOAD_FOLDER'], fn)
    f.save(fp)
    return jsonify({'success':True,'url':f'/uploads/{fn}','filename':fn})

# ─── ANSWER FILE IMPORT ───
@app.route('/api/import-answers', methods=['POST'])
@teacher_req
def api_import_answers():
    if 'file' not in request.files: return jsonify({'error':'No file'}),400
    f = request.files['file']
    content = ''
    fn = (f.filename or '').lower()
    if fn.endswith('.txt') or fn.endswith('.csv'):
        content = f.read().decode('utf-8', errors='ignore')
    elif fn.endswith('.docx'):
        import tempfile
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        f.save(tmp.name); tmp.close()
        try:
            from docx import Document
            doc = Document(tmp.name)
            content = '\n'.join([p.text for p in doc.paragraphs])
        except: content = ''
        os.unlink(tmp.name)
    else:
        content = f.read().decode('utf-8', errors='ignore')
    # Parse answers
    content = content.strip().upper()
    # Try format: 1.A 2.B 3.C or 1A 2B 3C
    import re
    answers = {}
    # Pattern: number followed by answer
    pairs = re.findall(r'(\d+)\s*[.)\-:]\s*([ABCD])', content)
    if pairs:
        for num, ans in pairs: answers[int(num)] = ans
    else:
        # Just a string of ABCD
        clean = re.sub(r'[^ABCDDS]', '', content)
        for i, ch in enumerate(clean): answers[i+1] = ch
    return jsonify({'success':True,'answers':answers,'total':len(answers),'raw':content[:500]})

# ─── DATA ───
RANKS = [(0,'Tân Binh','🐣','#9e9e9e','Mới bắt đầu hành trình!'),(10,'Đồng','🥉','#cd7f32','Chiến binh non trẻ!'),(20,'Bạc','🥈','#c0c0c0','Đang lên tay rồi đó!'),(30,'Vàng','🥇','#ffd700','Vàng 9999 luôn!'),(50,'Kim Cương','💎','#00bcd4','Sáng chói lóa mắt!'),(70,'Tinh Anh','🔮','#9c27b0','Cao thủ trong truyền thuyết!'),(90,'Cao Thủ','⚡','#ff5722','Thần đồng học tập!'),(150,'Ông Cố Nội','👑','#ff0000','Truyền thuyết sống!')]

def get_rank(total):
    rank = RANKS[0]
    for r in RANKS:
            if total >= r[0]: rank = r
    idx = RANKS.index(rank)
    nxt = RANKS[idx+1] if idx < len(RANKS)-1 else None
    prog = 0
    if nxt:
            needed = nxt[0] - rank[0]; done = total - rank[0]
            prog = min(100, int(done/needed*100)) if needed > 0 else 100
    else: prog = 100
    return {'name':rank[1],'icon':rank[2],'color':rank[3],'desc':rank[4],'min_subs':rank[0],
                'next_rank':{'name':nxt[1],'icon':nxt[2],'need':nxt[0]} if nxt else None,
                'progress':prog,'total_subs':total}

MASCOT_IDLE = ['Học đi bạn ơi, rank không tự lên đâu!','Cày 1 đề = 1 bước gần rank mới!','Hôm nay bạn đã làm bài chưa?','Ngày mai thi rồi, hôm nay phải cày!','AFK là thua, cày là thắng!','Não của bạn đang loading... hãy làm đề để upgrade!','Điểm 10 đang chờ bạn ở đề tiếp theo!','Cố công mài sắt, có ngày... lên rank Kim Cương!']
MASCOT_CORRECT = ['ĐỈNH QUÁ CHIẾN THẦN! 🔥','Não chạy bằng RTX 4090 rồi!','Câu này dễ mà, đúng không? 😎','Xuất sắc! Tiếp tục nhé!','GG! Too EZ!']
MASCOT_WRONG = ['Không sao, sai để học hỏi mà!','Câu này khó thật, đọc lại nhé!','Não cần sạc thêm cà phê rồi! ☕','Đừng buồn, còn nhiều câu dễ hơn!']
MASCOT_SUBMIT = {'god':['HỌC BÁ XUẤT HIỆN! 👑','Thần đồng là đây!'],'great':['Xuất sắc! Gần perfect rồi!'],'good':['Ổn áp! Chiến tiếp nhé! 💪'],'ok':['Được rồi, cần cố gắng thêm!'],'low':['Não cần sạc thêm cà phê! ☕','Calm down, học từ từ thôi!']}
RESULT_MEMES = {'god':[{'text':'HỌC BÁ XUẤT HIỆN!','emoji':'👑🔥'},{'text':'Full điểm! Thần đồng là bạn!','emoji':'🧠💎'}],'great':[{'text':'Gần perfect! Xuất sắc lắm!','emoji':'⭐🎯'},{'text':'Não RTX 4090 đã ON!','emoji':'🖥️🔥'}],'good':[{'text':'Ổn áp chiến tiếp!','emoji':'💪😎'}],'ok':[{'text':'Được rồi, cố gắng thêm!','emoji':'🤔💭'}],'low':[{'text':'Não cần sạc thêm cà phê!','emoji':'☕😴'},{'text':'Sai là để học, không phải để buồn!','emoji':'📚🌟'}]}

def read_file_content(fp):
    ext = fp.rsplit('.',1)[-1].lower() if '.' in fp else ''
    text = ''
    if ext == 'pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(fp)
                for page in reader.pages:
                    t = page.extract_text()
                    if t: text += t + '\n'
            except Exception as e: text = f'[PDF Error: {e}]'
    elif ext == 'docx':
            try:
                from docx import Document
                doc = Document(fp)
                for para in doc.paragraphs: text += para.text + '\n'
            except Exception as e: text = f'[DOCX Error: {e}]'
    elif ext in ('png','jpg','jpeg','gif','webp','bmp'):
            try:
                import pytesseract; from PIL import Image
                for tp in [r'C:\Program Files\Tesseract-OCR\tesseract.exe', r'C:\Users\LENOVO\AppData\Local\Tesseract-OCR\tesseract.exe']:
                    if os.path.exists(tp): pytesseract.pytesseract.tesseract_cmd = tp; break
                text = pytesseract.image_to_string(Image.open(fp), lang='vie+eng')
            except Exception as e: text = f'[OCR Error: {e}]'
    else:
            try:
                with open(fp, 'r', encoding='utf-8', errors='ignore') as f: text = f.read()
            except: pass
    return text.strip()

def parse_questions_smart(text):
    if not text or len(text) < 10: return []
    text = text.replace('\r\n','\n').replace('\r','\n')
    lines = text.split('\n')
    section = 'multiple_choice'
    section_map = {}
    for i, line in enumerate(lines):
            lu = line.strip().upper()
            if any(k in lu for k in ['TRAC NGHIEM NHIEU','PHUONG AN LUA CHON','NHIEU PHUONG AN','TRẮC NGHIỆM NHIỀU','PHƯƠNG ÁN LỰA CHỌN','NHIỀU PHƯƠNG ÁN']): section = 'multiple_choice'
            elif any(k in lu for k in ['DUNG SAI','ĐÚNG SAI','DUNG/SAI','ĐÚNG/SAI']): section = 'true_false'
            elif any(k in lu for k in ['TU LUAN','TỰ LUẬN']): section = 'essay'
            elif any(k in lu for k in ['TRA LOI NGAN','TRẢ LỜI NGẮN']): section = 'short_answer'
            section_map[i] = section
    q_positions = []
    for i, line in enumerate(lines):
            m = re.match(r'^\s*[Cc][aâ]u\s+(\d+)\s*[.:\)]\s*(.*)', line)
            if m: q_positions.append((i, int(m.group(1)), m.group(2).strip()))
    questions = []
    for idx, (li, qn, fl) in enumerate(q_positions):
            end = q_positions[idx+1][0] if idx+1 < len(q_positions) else len(lines)
            qt = section_map.get(li, 'multiple_choice')
            qlines = [fl] + [lines[j].strip() for j in range(li+1, end)]
            q = {'question_number': qn, 'content': '', 'type': qt, 'option_a':'','option_b':'','option_c':'','option_d':'','correct_answer':'','score':0.25}
            cl = []
            for line in qlines:
                line = line.strip()
                if not line: continue
                if qt == 'multiple_choice':
                    om = re.match(r'^\s*([ABCD])\s*[.)\s:]\s*(.+)', line)
                    if om: q[f'option_{om.group(1).lower()}'] = om.group(2).strip(); continue
                cl.append(line)
            q['content'] = '\n'.join(cl).strip() or '\n'.join(qlines).strip()[:300]
            if qt == 'true_false': q['score'] = 1.0
            elif qt == 'essay': q['score'] = 1.0
            if q['content'] and len(q['content']) > 2: questions.append(q)
    return questions

def ai_reply(msg):
    ml = msg.lower()
    if any(w in ml for w in ['nguyên hàm','tích phân','đạo hàm','integral','derivative']):
            return "Cô giáo AI: Về Nguyên hàm / Tích phân!\n\n📌 Công thức cơ bản:\n• ∫xⁿdx = xⁿ⁺¹/(n+1) + C\n• ∫sinx dx = −cosx + C\n• ∫cosx dx = sinx + C\n• ∫eˣdx = eˣ + C\n• ∫(1/x)dx = ln|x| + C\n\nPhương pháp: Đổi biến, tích phân từng phần, phân tích phân thức.\n\nEm muốn hỏi cụ thể phần nào?"
    if any(w in ml for w in ['mặt phẳng','mặt cầu','đường thẳng','tọa độ','oxyz']):
            return "Cô giáo AI: Hình học không gian Oxyz!\n\n📌 Công thức cần nhớ:\n• Mặt phẳng: ax+by+cz+d=0, VTPT n=(a,b,c)\n• Đường thẳng: (x−x₀)/a=(y−y₀)/b=(z−z₀)/c\n• Mặt cầu: (x−a)²+(y−b)²+(z−c)²=R²\n• d(M,mp) = |ax₀+by₀+cz₀+d| / √(a²+b²+c²)\n\nEm muốn hỏi phần nào cụ thể?"
    if any(w in ml for w in ['vật lý','newton','động lực','điện','từ trường','sóng']):
            return "Cô giáo AI: Môn Vật lý!\n\n📌 Công thức quan trọng:\n• F=ma, P=mg\n• Wd=mv²/2, Wt=mgh\n• U=IR, P=UI\n• T=2π√(l/g) (con lắc đơn)\n• v=λf (sóng cơ)\n\nEm cần giải bài tập cụ thể nào?"
    if any(w in ml for w in ['tin học','lập trình','python','sql','thuật toán']):
            return "Cô giáo AI: Môn Tin học!\n\n📌 Kiến thức trọng tâm:\n• Python: list, dict, loop, function\n• SQL: SELECT, WHERE, GROUP BY, JOIN\n• Thuật toán: tìm kiếm, sắp xếp\n• Độ phức tạp: O(1), O(n), O(n²), O(logn)\n\nEm muốn ôn phần nào?"
    if any(w in ml for w in ['xin chào','chào','hello','hi']):
            return "Chào em! 👋 Cô là AI Cô giáo của X-EXAM.\n\nCô có thể giúp em:\n• Giải thích kiến thức Toán, Lý, Tin\n• Gợi ý phương pháp làm bài\n• Phân tích điểm yếu cần cải thiện\n\nEm hỏi bất kỳ điều gì nhé! 😊"
    return f"Cô giáo AI nhận được câu hỏi: \"{msg[:80]}\"\n\nĐây là câu hỏi thú vị! Cô gợi ý:\n1. Xác định rõ kiến thức liên quan\n2. Viết ra công thức cần dùng\n3. Áp dụng từng bước cẩn thận\n4. Kiểm tra lại kết quả\n\nEm hỏi cụ thể hơn để cô hướng dẫn chi tiết nhé!"

# ─── SEED EXAM DATA ───
@app.route('/api/admin/seed-exams', methods=['POST'])
@admin_req
def api_seed_exams():
    c = get_db()
    admin = c.execute("SELECT id FROM users WHERE role='admin' ORDER BY id LIMIT 1").fetchone()
    admin_id = admin['id'] if admin else 1
    def mk(content, a='', b='', cc='', d='', ans='', t='mc'):
        return {'c':content,'a':a,'b':b,'cc':cc,'d':d,'ans':ans,'t':t}
    subjects = {
        'Toan':{'time':90,'qs':[
            mk('Ham so y=x^3-3x^2+2 co bao nhieu diem cuc tri?','0','1','2','3','C'),
            mk('Tap xac dinh y=log2(x-1):','(1;+inf)','[1;+inf)','(0;+inf)','(-inf;1)','A'),
            mk('Nguyen ham f(x)=2x+1:','x^2+x+C','x^2+C','2x^2+x+C','x+C','A'),
            mk('Tich phan 2f(x)dx (0->1) biet f=3:','3','6','9','12','B'),
            mk('z=3+4i co modun:','3','4','5','7','C'),
            mk('2^x=8, x=?','2','3','4','8','B'),
            mk('V hinh chop S.ABCD vuong canh a, SA=a:','a^3/3','a^3/6','a^3','a^3/2','A'),
            mk('KC tu O den MP x+2y-2z+3=0:','1','2','3','3/5','A'),
            mk('CSC u1=2,d=3. u10=?','27','29','30','32','B'),
            mk('log2(3)=a. log2(12)=?','a+2','2a','3a','a+1','A'),
            mk('y=(x+1)/(x-2) tiem can dung:','x=-1','x=1','x=2','x=-2','C'),
            mk('DT mat cau ban kinh R:','2piR2','4piR2','piR2','3piR2','B'),
            mk('a=(1;2;-1),b=(2;-1;3). a.b=?','-3','3','-1','1','A'),
            mk('x^2+y^2=9 ban kinh:','9','3','sqrt9','81','B'),
            mk('y=x^4-2x^2+1 dong bien:','(-1;0)','(0;1)','(-1;1)','(1;+inf)','D'),
            mk('Tich phan (2x+1)dx (0->2)=?','','','','','6','es'),
            mk('3^x=27. x=?','','','','','3','es'),
            mk('Hinh tru r=3,h=4. V=? (pi=3.14)','','','','','113.04','es'),
            mk('lim(3x^2+1)/(x^2-2) x->+inf=?','','','','','3','es'),
            mk('Lien hop z=2-3i la?','','','','','2+3i','es'),
        ]},
        'Vat ly':{'time':50,'qs':[
            mk('Don vi cuong do dong dien:','Volt','Ampe','Ohm','Watt','B'),
            mk('F=k*q1*q2/r^? (Coulomb)','1','2','3','4','B'),
            mk('Van toc song dien tu chan khong:','3x10^6','3x10^8','340','3x10^5','B'),
            mk('Quang dien la:','e bi but ra khi AS','Phat sang','Nhieu xa','Phan xa','A'),
            mk('f dao dong cuong buc:','f rieng','f bat ky','f luc cuong buc','0','C'),
            mk('T con lac don phu thuoc:','A,m','l,g','m,g','A,l','B'),
            mk('PU hat nhan toa nhiet co:','Do hut khoi +','Do hut khoi -','Khong HK','m tang','A'),
            mk('RLC cong huong khi:','R=0','ZL=ZC','ZL>ZC','ZL<ZC','B'),
            mk('AS buoc song ngan nhat:','Do','Vang','Luc','Tim','D'),
            mk('hf=A+Wd. A la:','Dong nang','Cong thoat','NL photon','The nang','B'),
            mk('x=5cos(2pi*t). Bien do:','2cm','5cm','10cm','2pi','B'),
            mk('Song co truyen trong:','Chan khong','MT vat chat','Ca hai','Ran','B'),
            mk('Tu cam khi:','I khong doi','I bien thien','I=0','q dung yen','B'),
            mk('Tia X buoc song:','Lon hon AS','Ngan hon AS','Bang AS','KXD','B'),
            mk('Luc Lorentz phuong:','Cung v','Cung B','Vuong goc v va B','Bat ky','C'),
            mk('m=2kg roi, g=10. v sau 3s=?','','','','','30','es'),
            mk('R=10,I=2. P=? W','','','','','40','es'),
            mk('Buoc song f=100MHz=? m','','','','','3','es'),
            mk('Con lac l=1,g=10. T=?','','','','','2.0','es'),
            mk('Wd electron v=10^6=? x10^-19 J','','','','','4.6','es'),
        ]},
        'Tieng Anh':{'time':60,'qs':[
            mk('She ___ to school every day.','go','goes','going','gone','B'),
            mk('If I ___ you, I would study harder.','am','was','were','be','C'),
            mk('The book ___ by millions.','has read','has been read','have been read','is reading','B'),
            mk('He is the man ___ car was stolen.','who','whom','whose','which','C'),
            mk('Different pronunciation: thought/though/through/thorough','thought','though','through','thorough','B'),
            mk('By the time we arrived, the movie ___.','started','has started','had started','was starting','C'),
            mk('___ it was raining, we went.','Because','Although','Since','When','B'),
            mk('Synonym of abundant:','rare','scarce','plentiful','limited','C'),
            mk('She suggested ___ to cinema.','go','to go','going','went','C'),
            mk('I wish I ___ speak French.','can','could','will','may','B'),
            mk('He has worked here ___ 2020.','for','since','in','from','B'),
            mk('Antonym of optimistic:','hopeful','confident','pessimistic','cheerful','C'),
            mk('Not until she arrived ___ about keys.','she realized','did she realize','she did','realized she','B'),
            mk('Children ___ play in park.','use to','used to','are used to','using to','B'),
            mk('People say he is genius = ?','He is said to be a genius.','He said to be.','He is say.','He was said be.','A'),
            mk('She started 5 years ago => has ___ for 5 years','','','','','been learning','es'),
            mk('If studied harder, he ___ exam','','','','','would have passed','es'),
            mk('Noun of succeed?','','','','','success','es'),
            mk('The more..the ___ (good)','','','','','better','es'),
            mk('Passive: built 1990 => ___ in 1990','','','','','was built','es'),
        ]},
        'Hoa hoc':{'time':50,'qs':[
            mk('Dien phan la qua trinh OxH xay ra khi dong dien:','2 chieu','1 chieu','Xoay chieu','Khong can','B'),
            mk('Anion trong DP nong chay di ve:','Anode bi OxH','Anode bi khu','Cathode bi khu','Cathode bi OxH','A'),
            mk('DP NaCl nong chay, o anode:','OxH Cl-','OxH Na+','Khu Cl-','Khu Na+','A'),
            mk('SP DP nong chay MgCl2:','MgO+HCl','Mg+H2+Cl2','MgCl2+H2O','Mg+Cl2','D'),
            mk('CT glucose:','C12H22O11','C6H12O6','(C6H10O5)n','C2H5OH','B'),
            mk('Polime thien nhien:','PE','PVC','Cao su buna','Tinh bot','D'),
            mk('Amino axit tinh:','Axit','Bazo','Luong tinh','Trung tinh','C'),
            mk('Amin bac 1:','CH3NH2','(CH3)2NH','(CH3)3N','C6H5NHCH3','A'),
            mk('DP NaCl mang ngan tao:','Na+Cl2','NaOH+Cl2+H2','NaCl+H2O','Na+H2','B'),
            mk('Fe+HCl:','FeCl3+H2','FeCl2+H2','Fe2O3+H2','FeCl2+H2O','B'),
            mk('Nuoc cung tam thoi:','Ca,Mg,Cl','Ca,Mg,HCO3','Na,K,SO4','Fe,Cu','B'),
            mk('Chat beo = este:','Glix+a.beo','Etanol+a.beo','Metanol+axit','Glix+HCl','A'),
            mk('Phenol tinh:','Axit yeu','Bazo manh','Trung tinh','Axit manh','A'),
            mk('Thuy phan saccarozo:','Glu','Fruc','Glu+Fruc','Tinh bot','C'),
            mk('Khi nha kinh:','O2','N2','CO2','H2','C'),
            mk('pH HCl 0.01M=?','','','','','2','es'),
            mk('CT muoi an?','','','','','NaCl','es'),
            mk('KL FeCl2 tu 5.6g Fe=?g','','','','','12.7','es'),
            mk('Mn oxi hoa trong KMnO4=?','','','','','+7','es'),
            mk('V H2 tu 2.3g Na+H2O=?lit','','','','','1.12','es'),
        ]},
        'Ngu van':{'time':120,'qs':[
            mk('Vo chong A Phu cua:','To Hoai','Kim Lan','NM Chau','N Tuan','A'),
            mk('Vi Hanh dung but phap:','Tuong trung','Chau','Nhan hoa','So sanh','B'),
            mk('Tay Tien tho cua:','Quang Dung','To Huu','Che Lan Vien','Huy Can','A'),
            mk('NV chinh Chiec thuyen ngoai xa:','Phung','Mien','My','Lien','A'),
            mk('Ai da dat ten cho dong song cua:','N Tuan','HPN Tuong','VT Phung','Nam Cao','B'),
            mk('PC Nguyen Tuan:','Gian di','Tai hoa uyen bac','Tru tinh','Hien thuc','B'),
            mk('Bai tho Song cua:','Xuan Quynh','HX Huong','Xuan Dieu','Huy Can','A'),
            mk('Nguoi lai do Song Da the loai:','Truyen ngan','Tieu thuyet','Tuy but','Kich','C'),
            mk('Ba cu Tu Vo nhat the hien:','Nhan hau','Phan khang','Kiem tinh','Me con','A'),
            mk('Hai dua tre doi lap:','AS-bong toi','Giau-ngheo','Tre-gia','Thi-que','A'),
            mk('PT bieu dat van nghi luan:','Tu su','Mieu ta','Nghi luan','Bieu cam','C'),
            mk('Dat Nuoc cua:','NK Diem','ND Thi','CL Vien','To Huu','A'),
            mk('Noi qua thuoc:','An du','Hoan du','Noi qua','Nhan hoa','C'),
            mk('Rung xa nu cua:','NT Thanh','NM Chau','NH Thiep','MV Khang','A'),
            mk('Hon Truong Ba cua:','NH Tuong','LQ Vu','VT Phung','ND Thi','B'),
            mk('Chu de Tay Tien?','','','','','Ve dep hao hoa lang man bi trang nguoi linh','es'),
            mk('Y nghia Chiec thuyen ngoai xa?','','','','','Nhin cuoc song tu nhieu goc do','es'),
            mk('3 tac pham VH 1945-1975?','','','','','Tay Tien, Viet Bac, Rung xa nu','es'),
            mk('Phep doi la gi?','','','','','Goi ten SV bang ten khac co QH gan gui','es'),
            mk('PC chinh luan HCM?','','','','','Lap luan chat che dan chung xac thuc','es'),
        ]},
        'Lich su':{'time':50,'qs':[
            mk('CM T8 thanh cong ngay:','2/9/1945','19/8/1945','1/5/1945','22/12/1944','B'),
            mk('Dien Bien Phu ngay:','30/4/1975','7/5/1954','2/9/1945','8/3/1965','B'),
            mk('Hiep dinh Paris nam:','1972','1973','1974','1975','B'),
            mk('Thong nhat dat nuoc:','2/9/1945','7/5/1954','30/4/1975','22/12/1976','C'),
            mk('LX phong ve tinh nam:','1955','1957','1961','1969','B'),
            mk('ASEAN thanh lap:','1965','1967','1970','1975','B'),
            mk('CTTG2 ket thuc:','1944','1945','1946','1947','B'),
            mk('Can Vuong do:','Ham Nghi+TTThuyet','PB Chau','PC Trinh','NAQ','A'),
            mk('DCSVN thanh lap:','3/2/1930','19/5/1930','2/9/1930','1/5/1930','A'),
            mk('Ap Bac 1963 o:','Dong Nai','My Tho','Ben Tre','Long An','B'),
            mk('Dong minh My:','Lien Xo','TQ','Anh','Cuba','C'),
            mk('Berlin sup:','1987','1989','1991','1993','B'),
            mk('CT lanh ket thuc:','LX tan ra 1991','Berlin 1989','My rut VN','TQ doi moi','A'),
            mk('Dong Du do:','PC Trinh','PB Chau','LV Can','HT Khang','B'),
            mk('LHQ thanh lap:','1944','1945','1946','1950','B'),
            mk('Y nghia CM T8?','','','','','Pha bo thong tri, lap VNDCCH','es'),
            mk('NN thang loi DBP?','','','','','Lanh dao dung dan, tinh than anh dung','es'),
            mk('5 nuoc ASEAN?','','','','','Thai Lan, Indonesia, Malaysia, Philippines, Singapore','es'),
            mk('ND Geneve 1954?','','','','','Cong nhan doc lap VN, chia vi tuyen 17','es'),
            mk('CTTG2 KT chau Au khi?','','','','','Duc dau hang 9/5/1945','es'),
        ]},
        'Sinh hoc':{'time':50,'qs':[
            mk('ADN co may mach:','1','2','3','4','B'),
            mk('Phien ma tao:','ADN','mARN','Protein','Lipit','B'),
            mk('Dot bien gen la:','Bien doi NST','Bien doi gen','So luong NST','Kieu hinh','B'),
            mk('Da boi NST:','2n','3n+','n','2n+1','B'),
            mk('CLTN tac dong:','Ca the','Quan the','Loai','HST','A'),
            mk('Pha sang quang hop o:','Chat nen','Tilacoit','Ti the','Nhan','B'),
            mk('Chuoi thuc an tu:','Tieu thu','San xuat','Phan giai','An thit','B'),
            mk('Aa la:','DH toi','DH tran','Di hop','KXD','C'),
            mk('PL doc lap cho:','Gen cung NST','Gen khac NST','Lien ket','Gioi tinh','B'),
            mk('Ho hap TB may GD:','2','3','4','5','B'),
            mk('Nguyen phan la:','Giam nhiem','Nguyen nhiem','TH protein','Phien ma','B'),
            mk('Ky sinh:','2 loi','1 loi 1 hai','2 hai','Trung tinh','B'),
            mk('Mang TB:','ADN','Lipit kep+protein','Tinh bot','Xenlulozo','B'),
            mk('Anticodon o:','mARN','tARN','rARN','snARN','B'),
            mk('HST gom:','QX+sinh canh','SV','MT','DV+TV','A'),
            mk('Aa x Aa ti le KH?','','','','','3:1','es'),
            mk('4 nucleotit ADN?','','','','','A, T, G, C','es'),
            mk('SP duong phan?','','','','','Axit piruvic','es'),
            mk('Bac DD cap 1?','','','','','Sinh vat san xuat','es'),
            mk('Nhan to tien hoa QT nho?','','','','','Bien dong di truyen','es'),
        ]},
        'Dia ly':{'time':50,'qs':[
            mk('VN o khu vuc:','Dong Bac A','Dong Nam A','Nam A','Tay A','B'),
            mk('DT VN khoang:','231k km2','331k km2','431k km2','131k km2','B'),
            mk('Song dai nhat VN:','Hong','Dong Nai','Cuu Long','Ma','C'),
            mk('DB lon nhat VN:','s.Hong','s.Cuu Long','Thanh Nghe','BT Thien','B'),
            mk('Khi hau VN:','On doi','Nhiet doi gio mua','Xich dao','Han doi','B'),
            mk('TP lon nhat VN:','HN','HP','DN','TP HCM','D'),
            mk('Vung KT nhat:','DB sHong','Dong Nam Bo','Tay Nguyen','Bac TB','B'),
            mk('KS Quang Ninh:','Dau mo','Than da','Sat','Vang','B'),
            mk('Rung nhiet doi o:','Nhiet thap','Mua lon nhiet cao','Kho','Gio','B'),
            mk('Dan so VN 2024:','80tr','90tr','100tr','110tr','C'),
            mk('Nganh GDP cao nhat:','NN','CN','DV','Khoang san','C'),
            mk('Bien Dong phia:','Bac','Tay','Dong','Nam','C'),
            mk('Tay Nguyen cay:','Lua','Ca phe','Mia','Bong','B'),
            mk('DB sHong dat:','Phu sa','Bazan','Cat','Phen','A'),
            mk('Nuoc DT lon nhat TG:','TQ','My','Canada','Nga','D'),
            mk('3 KS chinh VN?','','','','','Than, dau mo, sat','es'),
            mk('VN co ? tinh/TP','','','','','63','es'),
            mk('3 song lon VN?','','','','','Hong, Cuu Long, Dong Nai','es'),
            mk('Tinh DT lon nhat VN?','','','','','Nghe An','es'),
            mk('Do thi hoa la gi?','','','','','Tang ti le dan song thanh thi','es'),
        ]},
        'GDCD':{'time':50,'qs':[
            mk('Phap luat la:','QT do NN ban hanh','Dao duc','Tap quan','Quy uoc','A'),
            mk('Binh dang truoc PL:','Lam gi cung dc','CD binh dang Q va NV','Chi nu','Giau nhieu Q','B'),
            mk('Hien phap la:','Luat thuong','Luat co ban cao nhat','Nghi dinh','Thong tu','B'),
            mk('Q tu do KD:','HP','Luat DN','Ca HP+LDN','Khong co','C'),
            mk('Vi pham PL:','Trai DD','Trai PL co loi','Bat ky','Chi toi pham','B'),
            mk('Bau cu tu tuoi:','16','18','20','21','B'),
            mk('TNHS cho:','HC','Dan su','Toi pham','Ky luat','C'),
            mk('Binh dang gioi:','Giong nhau','DV co hoi nhu nhau','Nu nhieu Q','Nam nhieu Q','B'),
            mk('BK xam pham than the:','Ko bi bat trai PL','Lam gi cung dc','Ko lao dong','Ko thue','A'),
            mk('Dao duc la:','QT do PL','QT le phai dieu thien','Tap quan','Phong tuc','B'),
            mk('Bi mat thu tin thuoc:','KT','Tu do CN','CT','VH','B'),
            mk('16t bi TNHS khi:','Bat ky','Toi rat/DB NG','Ko bao gio','Toi nhe','B'),
            mk('Thuc hien PL may HT?','2','3','4','5','C'),
            mk('Quoc hoi co quan:','Hanh phap','Tu phap','Lap phap','Kiem sat','C'),
            mk('Muc dich PL:','Loi ich CN','Trat tu XH','Trung phat','Thue','B'),
            mk('4 HT thuc hien PL?','','','','','Tuan thu, Thi hanh, Su dung, Ap dung','es'),
            mk('Q tu do ngon luan?','','','','','Tu do bieu dat y kien','es'),
            mk('Phan biet DD va PL?','','','','','PL do NN cuong che, DD do du luan XH','es'),
            mk('2 quyen CD trong HP?','','','','','Hoc tap, lao dong','es'),
            mk('Tre em la duoi ? tuoi','','','','','16','es'),
        ]},
        'Tin hoc':{'time':50,'qs':[
            mk('Nhi phan may chu so?','8','10','2','16','C'),
            mk('1 Byte=? Bit','4','8','16','32','B'),
            mk('CPU la:','Central Process U','Central Processing Unit','Computer PU','Central Program U','B'),
            mk('Python in MH:','echo','printf','print()','write()','C'),
            mk('So nguyen Python:','float','str','int','bool','C'),
            mk('range(5) lap:','4','5','6','10','B'),
            mk('RAM la:','Chi doc','Truy cap NN mat khi tat','Vinh vien','Ko ghi','B'),
            mk('Sap xep don gian nhat:','Quick','Merge','Bubble','Heap','C'),
            mk('IPv4 = ? bit','16','32','64','128','B'),
            mk('SQL la:','Simple QL','Structured QL','System QL','Standard QL','B'),
            mk('SQL truy van:','INSERT','UPDATE','SELECT','DELETE','C'),
            mk('OS ma nguon mo:','Windows','macOS','Linux','iOS','C'),
            mk('Python list:','Bat bien','Mutable','Const','Ham','B'),
            mk('Binary search O(?):','O(n)','O(n^2)','O(logn)','O(1)','C'),
            mk('HTML la:','Ngon ngu LT','Danh dau SVB','QTCSDL','OS','B'),
            mk('13 = ? (nhi phan)','','','','','1101','es'),
            mk('len("Hello")=?','','','','','5','es'),
            mk('SELECT *..score>=8 nghia?','','','','','Lay HS diem tu 8','es'),
            mk('1GB=?MB','','','','','1024','es'),
            mk('[1,2,3].append(4)=?','','','','','[1, 2, 3, 4]','es'),
        ]},
    }
    ck2_hoa = [
        ('De on CK2 - Hoa 12: Dien phan (De 1)','Hoa hoc',50,[
            mk('DP la QT OxH khi dong dien nao di qua?','1 chieu','2 chieu','Xoay chieu','Khong can','A'),
            mk('Anion trong DP nong chay di ve:','Anode bi OxH','Anode bi khu','Cathode bi khu','Cathode bi OxH','A'),
            mk('Cation trong DP nong chay di ve:','Anode bi khu','Anode bi OxH','Cathode bi khu','Cathode bi OxH','C'),
            mk('DP NaCl nong chay, anode:','OxH Cl-','OxH Na+','Khu Cl-','Khu Na+','A'),
            mk('SP DP nong chay MgCl2:','MgO+HCl','Mg+H2+Cl2','MgCl2+H2O','Mg+Cl2','D'),
            mk('DP NaCl nong chay+dd NaCl deu co o anode:','O2','Na','Cl2','H2','C'),
            mk('DP CaCl2 nong chay, cathode:','OxH Ca2+','Khu Ca2+','OxH Cl-','Khu Cl-','B'),
            mk('DP NaOH nong chay, anode:','Na++e->Na','4OH-->O2+2H2O+4e','2OH-->H2+O2+2e','2O2-->O2+4e','B'),
            mk('Thu tu DP ion KL o cathode:','The am hon DP truoc','The duong hon DP truoc','The duong o anode','The am o anode','B'),
            mk('DP dd CuSO4 dien cuc tro, anode:','Cu2++2e->Cu','2H2O->4H++O2+4e','Cu->Cu2++2e','2H2O+2e->H2+2OH-','B'),
            mk('DP dd NaCl mang ngan, anode+cathode:','Cl2 va NaOH+H2','Na va Cl2','Cl2 va Na','NaOH va H2','A'),
            mk('DP dd CuSO4, ion DP dau tien o cathode:','Cu2+','H+ nuoc','SO42-','OH- nuoc','A'),
            mk('DP dd NaCl, o cathode:','OxH Na+','OxH H2O','Khu H2O','Khu Na+','C'),
            mk('DP Cu(NO3)2 dien cuc tro, anode:','H2O->O2+2H++2e','2H2O+2e->H2+2OH-','Cu->Cu2++2e','Cu2++2e->Cu','A'),
            mk('Ion KL bi DP trong dd (graphite):','Na+','Cu2+','Ca2+','K+','B'),
            mk('DP dd AgNO3, cathode:','','','','','Ag++e->Ag','es'),
            mk('DP dd CuSO4 mau xanh nhat dan vi:','','','','','Cu2+ bi khu thanh Cu','es'),
            mk('PU o cathode khi DP Al2O3 nong chay:','','','','','Al3++3e->Al','es'),
            mk('DP dd NaCl+NaBr, anode OxH dau tien:','','','','','2Br-->Br2+2e','es'),
            mk('Thu tu DP cathode dd Fe2+,Fe3+,Cu2+:','','','','','Fe3+, Cu2+, Fe2+','es'),
        ]),
        ('De on CK2 - Hoa 12: Kim loai (De 2)','Hoa hoc',50,[
            mk('Cau hinh e Mg(Z=12) lop ngoai:','3s1','3s2','3s2p1','3p1','B'),
            mk('Nguyen to co e lop ngoai cung it nhat:','C','Li','O','F','B'),
            mk('Nguyen to co 3e lop ngoai:','11Na','13Al','20Ca','26Fe','B'),
            mk('Khong co trong mang tinh the KL:','Ion KL','Electron','Nguyen tu KL','Anion goc acid','D'),
            mk('KL nhiet do nc thap nhat:','Au','Cu','Na','Hg','D'),
            mk('KL do cung lon nhat:','Cr','Al','Mg','Na','A'),
            mk('KL nhiet do nc cao nhat:','Na','Pb','Hg','W','D'),
            mk('KL KLR nho nhat:','Na','Li','Fe','Al','B'),
            mk('Electron tu do trong mang KL:','Theo QD XD','Xung quanh vi tri','Tu do toan mang','Khu vuc nhat dinh','C'),
            mk('Lien ket trong mang KL:','Ion','Cong hoa tri','Van der Waals','Kim loai','D'),
            mk('Au dat mong keo soi vi:','Tinh deo cao','Tinh dan dien','Do cung cao','Nhiet do nc cao','A'),
            mk('KL trang thai long o dk thuong:','Na','Hg','Al','Fe','B'),
            mk('Cu,Al lam day dan dien vi:','Nhiet do nc cao','Tinh dan nhiet','Tinh dan dien','Anh kim','C'),
            mk('W lam day toc vi:','Nhiet do nc cao','Dan dien tot','Tinh deo','Do cung','A'),
            mk('Tinh chat VL chung KL:','Deo,dan dien,nc cao','Deo,dan dien nhiet,anh kim','Dan dien nhiet,KLR lon,anh kim','Deo,anh kim,rat cung','B'),
            mk('Thu tu dan dien giam dan:','','','','','Ag>Cu>Au>Fe (Ag tot nhat)','es'),
            mk('Nguyen tac dieu che KL:','','','','','Khu ion KL thanh nguyen tu','es'),
            mk('Day KL DP hop chat nong chay:','','','','','Na, Ca, Al','es'),
            mk('Oxide bi H2 khu tao KL:','','','','','CuO, Fe2O3, PbO...','es'),
            mk('Phuong phap dieu che Mg tu MgCl2:','','','','','Dien phan MgCl2 nong chay','es'),
        ]),
    ]
    ck2_ls = [
        ('De on CK2 - Lich su 12: Doi ngoai + HCM (De 1)','Lich su',50,[
            mk('Muc dich doi ngoai VN dau TK XX:','Tranh thu ung ho quoc te','Mo rong lanh tho','PT kinh te','Hop tac quan su','A'),
            mk('NAQ sang Phap nam:','1911','1920','1930','1945','A'),
            mk('NAQ sang lap o Phap:','Dang CS Phap','Dang XH Phap','Quoc te CS','Hoi VNCMTN','A'),
            mk('Hiep dinh Geneve 1954 cong nhan:','Doc lap chu quyen VN','VN thong nhat','VN thuoc dia','My rut quan','A'),
            mk('VN gia nhap ASEAN nam:','1967','1975','1995','2000','C'),
            mk('VN gia nhap LHQ nam:','1945','1975','1977','1986','C'),
            mk('Doi ngoai VN tu 1986:','Da phuong hoa da dang hoa','Chi hop tac XHCN','Dong cua KT','Chi hop tac My','A'),
            mk('VN gia nhap WTO nam:','2000','2004','2007','2010','C'),
            mk('Que HCM:','Ha Tinh','Nghe An','Thanh Hoa','Hue','B'),
            mk('HCM ra di tim duong cuu nuoc nam:','1905','1911','1920','1930','B'),
            mk('HCM doc TN doc lap ngay:','19/8/1945','2/9/1945','22/12/1946','7/5/1954','B'),
            mk('HCM sang lap DCSVN ngay:','3/2/1930','19/5/1930','2/9/1945','22/12/1944','A'),
            mk('UNESCO vinh danh HCM la:','Anh hung GPDT nha VH kiet xuat','Nha khoa hoc','Nha kinh te','Nha quan su','A'),
            mk('HCM mat nam:','1965','1968','1969','1975','C'),
            mk('Con duong cuu nuoc HCM chon:','CM vo san','CM tu san','Cai cach','Bao luc','A'),
            mk('Con duong cuu nuoc HCM?','','','','','Doc lap dan toc gan voi CNXH','es'),
            mk('Y nghia thanh lap DCSVN 1930?','','','','','Cham dut khung hoang duong loi lanh dao CM VN','es'),
            mk('Vai tro HCM trong CM T8?','','','','','Lanh dao toan dan tong KN gianh chinh quyen','es'),
            mk('TN doc lap khang dinh gi?','','','','','Quyen doc lap tu do cua dan toc VN','es'),
            mk('Bai hoc tu HCM?','','','','','Hoc tap suot doi yeu nuoc phuc vu nhan dan','es'),
        ]),
        ('De on CK2 - Lich su 12: HCM anh hung GPDT (De 2)','Lich su',50,[
            mk('HCM gui Den Versailles yeu cau (1919):','Cac quyen tu do cho nhan dan VN','Doc lap hoan toan','Lien bang Dong Duong','Rut quan Phap','A'),
            mk('HCM doc Luan cuong Lenin nam:','1917','1920','1925','1930','B'),
            mk('Hoi VNCMTN sang lap nam:','1920','1925','1930','1941','B'),
            mk('HCM la Chu tich nuoc tu nam:','1930','1941','1945','1946','C'),
            mk('Duong loi KC chong Phap do HCM:','Toan dan toan dien lau dai tu luc','Danh nhanh','Chi vien tro','Thuong luong','A'),
            mk('HCM gui thu HS 9/1945:','Non song VN co tro nen ve vang...','Hoc lam viec','Hoc thi','Hoc kinh doanh','A'),
            mk('Di chuc HCM ve:','Doan ket trong Dang va nhan dan','Chien tranh','Kinh te','Ngoai giao','A'),
            mk('HCM thanh lap Mat tran VM nam:','1930','1936','1941','1945','C'),
            mk('HCM viet Ban an che do thuc dan Phap nam:','1920','1925','1930','1935','B'),
            mk('Phan tich duong loi KC chong Phap:','Toan dan toan dien truong ky','Chi quan su','Chi CT','Chi KT','A'),
            mk('Vai tro HCM trong DCSVN:','Trieu tap hoi nghi hop nhat 3 to chuc','Chi dao tu xa','Khong truc tiep','Lanh dao tu nuoc ngoai','A'),
            mk('So do luoi tan VN gia nhap:','1945','1975','1977','2000','C'),
            mk('DL doi ngoai VN hien nay:','Tiep tuc mo cua hoi nhap','Dong cua','Chi My','Chi XHCN','A'),
            mk('FTA VN ky ket dau la:','ASEAN FTA','VKFTA (Han Quoc)','EVFTA (EU)','CPTPP','A'),
            mk('Thanh tuu doi ngoai VN sau 1975:','Gia nhap LHQ ASEAN WTO','Chi ASEAN','Chi LHQ','Chi WTO','A'),
            mk('Phan tich vai tro HCM voi CM VN?','','','','','Tim ra con duong cuu nuoc dung dan lanh dao CM thanh cong','es'),
            mk('Tai sao HCM chon CM vo san?','','','','','Chi CM vo san giai phong triet de DT bi ap buc','es'),
            mk('Y nghia HCM doc TN doc lap?','','','','','Khai sinh nuoc VNDCCH khang dinh quyen doc lap','es'),
            mk('Bai hoc HCM ve tu hoc?','','','','','Tu hoc suot doi hoc o moi noi moi luc','es'),
            mk('Trach nhiem ban than tu cuoc doi HCM?','','','','','Hoc tap lam viec cong hien cho DT','es'),
        ]),
    ]
    ck2_tin = [
        ('De on CK2 - Tin hoc 12: Hoi thao + Ket noi (De 1)','Tin hoc',45,[
            mk('Muc tieu chinh hoi thao HN:','Gioi thieu CN moi','Giup HS chon nghe','Giai tri','The thao','B'),
            mk('Yeu to quan trong nhat lap KH HT:','Chon ngay gio','Dia diem','Xac dinh doi tuong','Moi dien gia','C'),
            mk('Yeu to quan trong nhat bai trinh bay:','Hinh anh sinh dong','ND ro rang de hieu','Thoi luong','Phong chu','B'),
            mk('Xac dinh dau tien khi lap KH HT:','Chu de HT','DS dien gia','Chi phi','So nguoi','A'),
            mk('Nguoi dieu phoi HT nen:','Gioi thieu DG','DK TG lich trinh','Tuong tac KG','Tao khong khi','B'),
            mk('PP danh gia ket qua HT:','Khao sat phan hoi','So luong nguoi','Thoi luong','So cau hoi','A'),
            mk('Cong KN thiet bi hien thi voi may tinh:','USB','HDMI','Bluetooth','Ethernet','B'),
            mk('KN BT may tinh-DT, can truoc:','Ghep doi','Bat BT ca hai','Xac nhan KN','Truyen file','B'),
            mk('Nha thong minh DK qua:','Cap tin hieu','Wi-Fi','BT va Wi-Fi','Chi cap','C'),
            mk('BT thuong dung de:','KN khoang cach xa','Chia se DL ngan','KN Internet','Chi may in','B'),
            mk('BT thuoc loai KN:','Co day','Ket noi xa','Ket noi gan','Cap tin hieu','C'),
            mk('Dac diem nha thong minh:','Hoan toan thu cong','Tu dong hoa DK tu xa','Khong can Internet','Chi HDMI','B'),
            mk('Cap Ethernet dung:','KN man hinh','Truyen DL mang LAN','KN luu tru','Truyen am thanh','B'),
            mk('Phan chan trang chua:','ND chinh','Ban quyen MXH','Dau trang','Thong bao','B'),
            mk('Phan dau trang nhu trang bia:','Than trang','Dau trang','Chan trang','Thong bao','B'),
            mk('Favicon la gi?','','','','','Bieu tuong dai dien trang web tren tab trinh duyet','es'),
            mk('Google Sites la gi?','','','','','Phan mem Google ho tro xay dung trang web mien phi','es'),
            mk('Buoc dau tien XD trang web?','','','','','Dinh hinh y tuong: xac dinh muc dich va doi tuong','es'),
            mk('Menu dropdown la gi?','','','','','Menu tha xuong khi nguoi dung hover hoac click vao','es'),
            mk('DK den thong minh khi khong o nha, dung?','','','','','Wi-Fi ket noi Internet (IoT)','es'),
        ]),
        ('De on CK2 - Tin hoc 12: Thiet ke trang web (De 2)','Tin hoc',45,[
            mk('De dinh hinh y tuong trang web can:','Xay dung kien truc ND','Lam logo','Chon phan mem','Xac dinh muc dich doi tuong','D'),
            mk('Thiet ke my thuat trang web gom:','Chi chon mau','Bang mau phong chu bieu tuong','Chi lam logo','Chi chon anh','B'),
            mk('Phan than trang web chua:','ND chinh','Chi logo','Chi menu','Chi lien ket','A'),
            mk('Thanh dieu huong thuong o:','Chan trang','Than trang','Dau trang','Khong co','C'),
            mk('Thiet ke header, dieu quan trong:','Dat toan bo ND vao dau','Sap xep ro rang de su dung','Them nhieu HU','Chi 1 mau','B'),
            mk('Thanh tim kiem thuong o:','Chan trang','Than trang','Dau trang','Khong can','C'),
            mk('Chuan bi tu lieu trang web:','Chi lam 1 lan','Keo dai suot du an','Khong can thiet','Chi o cuoi','B'),
            mk('Kieu chu dam thuong de:','Trang tri','Phan biet tu ngu','Lam dep','Khong y nghia','B'),
            mk('Nha thong minh thiet bi:','Chi co day','KN voi nhau qua mang','Hoan toan doc lap','Khong can dien','B'),
            mk('HDMI truyen:','DL mang','Hinh anh am thanh','Tin hieu BT','File van ban','B'),
            mk('WiFi la KN:','Co day','Khong day','Chi noi bo','Chi ngoai troi','B'),
            mk('Chieu PPT tu laptop len TV, day hong, dung:','HDMI','USB','Wi-Fi Miracast','Ethernet','C'),
            mk('Moi trang web day du gom may phan:','2','3','4','5','B'),
            mk('Chuyen bao thi khong nen:','Chuyen mau','Doi phong chu','Them nhieu hieu ung phuc tap','Chinh bo cuc','C'),
            mk('Khi chon bang mau trang web:','Tuong phan manh','1 mau duy nhat','Tuong phan nhe nhang hai hoa','Mau ngau nhien','C'),
            mk('Quy trinh XD trang web gom may buoc?','','','','','4 buoc: Xac dinh - Dinh hinh - Thiet ke - Xay dung - Kiem tra','es'),
            mk('Su khac biet giua website va trang web?','','','','','Website la tap hop nhieu trang web lien ket voi nhau','es'),
            mk('Hover effect la gi?','','','','','Hieu ung hien thi khi di chuot qua phan tu trang web','es'),
            mk('Site map la gi?','','','','','So do the hien cau truc va to chuc noi dung cua website','es'),
            mk('Bao nhieu muc trong menu la hop ly?','','','','','5-7 muc chinh de nguoi dung de nho va su dung','es'),
        ]),
    ]
    created = 0
    total_q = 0
    all_exams = []
    for subj, data in subjects.items():
        all_exams.append(('De on thi TN THPT - ' + subj + ' 2026', subj, data['time'], data['qs']))
    for title, subj, time, qs in ck2_hoa + ck2_ls + ck2_tin:
        all_exams.append((title, subj, time, qs))
    for title, subject, time_limit, questions in all_exams:
        existing = c.execute("SELECT id FROM exams WHERE title=?", (title,)).fetchone()
        if existing:
            print("SKIP: " + title)
            continue
        cur = c.execute("INSERT INTO exams(title,subject,description,time_limit,total_score,created_by,is_open,max_attempts,shuffle_questions,teacher_approved) VALUES(?,?,?,?,?,?,?,?,?,?)",
            (title, subject, "De on tap chinh thuc - Truong THPT Duy Tan - Co dap an", time_limit, 10, admin_id, 1, 999, 1, 1))
        eid = cur.lastrowid
        for i, q in enumerate(questions, 1):
            qtype = 'multiple_choice' if q['t'] == 'mc' else 'essay'
            score = 0.5 if qtype == 'multiple_choice' else 1.0
            c.execute("INSERT INTO questions(exam_id,question_number,type,content,option_a,option_b,option_c,option_d,correct_answer,score) VALUES(?,?,?,?,?,?,?,?,?,?)",
                (eid, i, qtype, q['c'], q.get('a',''), q.get('b',''), q.get('cc',''), q.get('d',''), q.get('ans',''), score))
        created += 1
        total_q += len(questions)
    c.commit()
    c.close()
    return 'OK: Created ' + str(created) + ' exams with ' + str(total_q) + ' questions!'

# Auto-init DB on import (for gunicorn/render)
init_db()

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(os.path.dirname(DB), exist_ok=True)
    init_db()
    print()
    print("="*55)
    print("  X-EXAM v5.0 - Hệ thống luyện thi THPT")
    print("  Mở trình duyệt: http://localhost:5000")
    print("  Tài khoản: admin/admin123 | student/123456")
    print("="*55)
    print()
    app.run(debug=False, host='0.0.0.0', port=5000, threaded=True)
